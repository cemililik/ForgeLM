"""Document ingestion — turn raw files (PDF/DOCX/EPUB/TXT) into SFT-ready JSONL.

Phase 11 (Document Ingestion) — bridges raw enterprise corpora (legal, medical,
policy manuals) to ForgeLM's training data format without forcing operators to
write custom preprocessing.

Output shape is the ``{"text": "<chunk>"}`` JSONL ForgeLM's data loader already
recognizes as pre-formatted SFT input (see ``forgelm/data.py``). Pair with
``forgelm --generate-data`` if you want the chunks expanded into Q&A
``messages`` form via a teacher model.

Optional extra:

    pip install 'forgelm[ingestion]'

OCR is out of scope. PDFs without a text layer produce a warning and zero
chunks; pre-process with Tesseract / AWS Textract before ingest.

Public API:

* :class:`IngestionResult` — outcome dataclass
* :func:`ingest_path` — single file or directory → JSONL
* :func:`list_supported_formats` — informational helper for the CLI
* :func:`describe_strategies` — chunking strategy descriptions
"""

from __future__ import annotations

import json
import logging
import math
import os
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Tuple

from ._pypdf_normalise import DEFAULT_PROFILE as _DEFAULT_NORMALISE_PROFILE
from ._pypdf_normalise import apply_profile as _apply_normalise_profile
from ._pypdf_normalise import count_substitutions as _count_normalise_substitutions
from ._script_sanity import (
    DEFAULT_THRESHOLD as _DEFAULT_SCRIPT_SANITY_THRESHOLD,
)
from ._script_sanity import (
    ScriptSanityReport,
    check_script_sanity,
)
from ._strip_pattern import apply_strip_patterns as _apply_strip_patterns

logger = logging.getLogger("forgelm.ingestion")


# ---------------------------------------------------------------------------
# Public types
# ---------------------------------------------------------------------------


SUPPORTED_EXTENSIONS: Tuple[str, ...] = (".pdf", ".docx", ".epub", ".txt", ".md")


class OptionalDependencyError(ImportError):
    """Raised when an optional ingestion extra (PDF / DOCX / EPUB) is missing.

    Subclasses :class:`ImportError` so existing call sites that catch the
    broader class keep working, while CLI dispatchers can opt into the narrower
    type to distinguish operator-actionable "install the extra" failures from
    genuine import bugs inside ``forgelm`` itself (which should propagate with
    their original traceback rather than be swallowed and re-emitted as a
    generic install hint).
    """


class IngestParameterError(ValueError):
    """Raised when an operator-supplied parameter is invalid against a real file.

    Phase 15 introduced page-range / strip-pattern / page-range-vs-page-count
    validators that run *inside* the per-file extractor — the existing
    soft-fail catch in :func:`_extract_text_for_ingest` would otherwise log
    the failure and continue with the rest of the corpus, hiding the
    operator's mistake. ``IngestParameterError`` propagates through the
    soft-fail catch so the CLI dispatcher can translate it to
    ``EXIT_CONFIG_ERROR`` per the contract.

    Subclasses :class:`ValueError` so existing call sites (and the CLI's
    documented exception-handling shape) keep working without an explicit
    ``except`` for this subclass.
    """


CHUNK_STRATEGIES: Tuple[str, ...] = ("sliding", "paragraph", "markdown", "semantic")

# Validation messages — pinned as module constants so ruff / SonarCloud
# don't flag the duplicated string literal across the four chunkers below.
_CHUNK_TOKENS_POSITIVE_MSG: str = "chunk_tokens must be positive"
_CHUNK_SIZE_POSITIVE_MSG: str = "max_chunk_size must be positive"
_MARKDOWN_OVERLAP_UNSUPPORTED_MSG: str = (
    "overlap / overlap_tokens is not supported for --strategy markdown "
    "(sections are atomic; an overlap would slice mid-section and break "
    "the heading-breadcrumb invariant). Use --strategy sliding for "
    "overlapping windows."
)


@dataclass
class IngestionResult:
    """Summary of an ``ingest_path`` run.

    Both :attr:`extra_notes` (free-text, operator-friendly) and
    :attr:`notes_structured` (programmatic ``{key: value}``) are emitted
    so machine-driven pipelines do not need to regex-match the prose. The
    structured payload is stable across releases; the free-text list is
    for human consumption and may rephrase the same facts.

    Phase 15 added:

    * ``pdf_paragraph_packed_lines_stripped`` — second-pass header dedup
      count surfaced for parity with the page-level
      :attr:`pdf_header_footer_lines_stripped`.
    * ``script_sanity_triggered`` / ``strip_pattern_substitutions`` /
      ``urls_handled`` / ``frontmatter_pages_dropped`` —
      operator-facing counts for the new Wave 1 / Wave 2 behaviours.
    """

    output_path: Path
    chunk_count: int
    files_processed: int
    files_skipped: int
    total_chars: int
    format_counts: Dict[str, int] = field(default_factory=dict)
    pii_redaction_counts: Dict[str, int] = field(default_factory=dict)
    secrets_redaction_counts: Dict[str, int] = field(default_factory=dict)
    extra_notes: List[str] = field(default_factory=list)
    notes_structured: Dict[str, Any] = field(default_factory=dict)
    pdf_header_footer_lines_stripped: int = 0
    pdf_paragraph_packed_lines_stripped: int = 0
    script_sanity_triggered: int = 0
    strip_pattern_substitutions: int = 0
    urls_handled: int = 0
    frontmatter_pages_dropped: int = 0


# ---------------------------------------------------------------------------
# Phase 15 — extraction context + post-extract normalisations
# ---------------------------------------------------------------------------


# Default skip-list for EPUB items the operator almost never wants to
# train on (TOC, cover, copyright, colophon, title-page, front-matter).
# Implemented as a frozenset of normalised lower-case substrings so a
# file named ``cover.xhtml`` or an ``epub:type="cover"`` declaration both
# match without per-source spelling drift.
_DEFAULT_EPUB_SKIP_ITEMS: Tuple[str, ...] = (
    "nav",
    "toc",
    "cover",
    "copyright",
    "colophon",
    "titlepage",
    "frontmatter",
)


# ``[ \t]*`` not ``\s*``: callers strip linewise so the ``\s`` form
# would over-match newlines and tab into the YAML body. Anchored on
# both sides so it stays a state-machine-friendly probe rather than
# the regex we apply to the body itself.
_YAML_FRONTMATTER_PATTERN = re.compile(r"\A---[ \t]*\n.*?\n---[ \t]*\n", re.DOTALL)


# Operator-facing URL pattern for ``--strip-urls`` (Phase 15 Task 14).
# The character class is bounded and non-overlapping, so an unbounded
# ``+`` does not backtrack catastrophically — there is no greedy/lazy
# interaction across a shared character set. Dropping the original
# ``{1,2048}`` upper bound (review round-2 S-5) keeps a 3-KB URL from
# being truncated mid-string and left as a partial residue in strip
# mode; URL handling now consumes the URL to its natural whitespace /
# punctuation boundary in one match.  ``A-Z`` only (no ``a-z``) under
# IGNORECASE — Sonar python:S5869 false-positive otherwise.
_URL_PATTERN = re.compile(r"\b(?:https?|ftp)://[A-Z0-9._~:/?#\[\]@!$&'()*+,;=%-]+", re.IGNORECASE)


@dataclass
class _ExtractContext:
    """Mutable context threaded through the per-file extraction path.

    Bundles every Phase 15 knob so individual extractors do not balloon
    in argument count. Mutable in three places only:

    * ``dedup_state`` — counters accumulated from PDF header / footer
      dedup (existing pre-Phase-15 behaviour).
    * ``script_sanity_reports`` — list of :class:`ScriptSanityReport`
      objects produced by every per-file sanity check (one entry per
      file, triggered or not).
    * ``frontmatter_dropped_pages`` — indices of PDF front-matter pages
      the heuristic dropped, for inclusion in the structured notes.

    Everything else is read-only after construction.
    """

    dedup_state: Optional[Dict[str, int]] = None
    page_range: Optional[Tuple[int, int]] = None  # 1-indexed inclusive
    normalise_profile: str = _DEFAULT_NORMALISE_PROFILE
    language_hint: Optional[str] = None
    script_sanity_threshold: float = _DEFAULT_SCRIPT_SANITY_THRESHOLD
    script_sanity_reports: List[ScriptSanityReport] = field(default_factory=list)
    keep_md_frontmatter: bool = False
    epub_skip_frontmatter: bool = True
    epub_skip_items: Tuple[str, ...] = _DEFAULT_EPUB_SKIP_ITEMS
    # Phase 15 Wave 2 Task 13: front-matter / back-matter heuristic is DEFAULT ON
    # in v0.6.0 (audit recommended; opt-out via --keep-frontmatter to preserve
    # the pre-Phase-15 "keep every page" behaviour). The library-level default
    # also moves from True to False so a library caller without the kwarg sees
    # the new behaviour — operators who freeze a script + .yaml + .py can
    # restore the pre-15 behaviour with one explicit kwarg.
    keep_frontmatter: bool = False
    strip_patterns: List[Tuple[str, "re.Pattern[str]"]] = field(default_factory=list)
    strip_pattern_timeout: Optional[int] = 5
    strip_urls_mode: str = "keep"  # keep | mask | strip
    frontmatter_dropped_pages: List[int] = field(default_factory=list)
    urls_handled_total: int = 0
    strip_pattern_substitutions_total: int = 0


# ---------------------------------------------------------------------------
# Format extractors — each returns the document's plain text or raises
# ---------------------------------------------------------------------------


_PDF_REPEAT_MIN_PAGES: int = 3
_PDF_REPEAT_THRESHOLD: float = 0.7
_PDF_EDGE_WINDOW: int = 3
"""Phase 15 Task 1 — number of rows from each page edge inspected per pass.

The pre-Phase-15 implementation only considered the **topmost** / **bottommost**
line per page on each pass. That worked for the homogeneous single-line
header case but silently exited the dedup loop the moment the outermost
row varied per page (e.g. a per-section title on top, a page number on
the bottom): no top/bottom recurrence at the cutoff → exit pass 1 →
never reach the constant publication-identifier line sitting one row
deeper. The audit (§1.1, 2026-05-11) traced 74/82 of a Turkish PDF's
chunks back to exactly this trap.

Widening the inspection to the top-3 / bottom-3 rows per page lets a
line recurring at **any** position within that window get stripped in
the same pass, so a variable-outer-line corpus no longer locks the
deeper-row constant out. Three is conservative: real-world page edges
rarely carry more than a 2-line running header + a 1-line page number
or copyright row, so 3 catches the modal case while leaving paragraph
text alone (a paragraph that legitimately re-appears at row 4+ across
≥ 70 % of pages is itself almost-certainly boilerplate worth flagging,
but we do not strip it because that is outside the documented header
/ footer contract).
"""


def _windowed_repeating_edge_lines(
    page_lines: List[List[str]],
    cutoff: int,
    window: int = _PDF_EDGE_WINDOW,
) -> Tuple[set, set]:
    """Return the set of lines recurring near each page's top / bottom edge.

    For each page, inspects up to ``window`` lines from the start and
    from the end (de-duplicated per page so a 1-line page does not
    inflate the count). A line counted at any offset within the window
    is treated as "edge-class" — when it recurs in ≥ ``cutoff`` pages it
    is eligible to be stripped, no matter whether it lived at the
    absolute outermost slot on every page or shifted by one row when
    the outer slot was occupied by a variable-content line.
    """
    first_counts: Counter = Counter()
    last_counts: Counter = Counter()
    for lines in page_lines:
        if not lines:
            continue
        # ``set`` de-dupes inside a single page so a short page that
        # happens to repeat a line wouldn't drive up the per-page count
        # past the integer 1.
        first_counts.update(set(lines[:window]))
        last_counts.update(set(lines[-window:]))
    return (
        {ln for ln, n in first_counts.items() if n >= cutoff},
        {ln for ln, n in last_counts.items() if n >= cutoff},
    )


def _pop_one_page_window(
    lines: List[str],
    repeating_firsts: set,
    repeating_lasts: set,
    window: int,
) -> Tuple[List[str], int]:
    """Per-page worker for :func:`_pop_windowed_edges`.

    Round-3 review (CodeRabbit) — the head walk inspects every
    position in the top-N window and drops any line whose stripped
    form is in ``repeating_firsts``, but the pre-round-3 tail walk
    peeled consecutively from the very end and stopped at the first
    non-match. That asymmetry let a *constant deeper bottom-edge*
    line survive when the outermost last line varied — the audit
    §1.1 trap re-occurring at the bottom of the page (variable page
    number on the last line, constant footer one row deeper):

    .. code-block:: text

       page N: [..., "Body N.", "FOOTER", "N"]

    With the outermost line varying ("1" / "2" / ...) the asymmetric
    tail walk would never reach the constant ``"FOOTER"`` one row
    deeper. The symmetric walk below mirrors the head walk so a
    bottom-edge constant line at any offset within the window is
    stripped in a single pass.
    """
    kept: List[str] = []
    dropped = 0
    for idx, ln in enumerate(lines):
        if idx < window and ln in repeating_firsts:
            dropped += 1
            continue
        kept.append(ln)
    # Symmetric tail walk: drop any line within the last ``window``
    # positions whose stripped form is in ``repeating_lasts``. Walking
    # in reverse so ``offset`` mirrors the head walk's ``idx``.
    surviving: List[str] = []
    total = len(kept)
    for rev_idx, ln in enumerate(reversed(kept)):
        # rev_idx == 0 is the last line; rev_idx < window covers the
        # bottom-N window inclusive.
        if rev_idx < window and ln in repeating_lasts:
            dropped += 1
            continue
        surviving.append(ln)
    # ``surviving`` is in reversed order; flip to restore page order.
    kept = list(reversed(surviving))
    # Suppress unused-name warning on ``total`` — kept as documentation
    # of the window's reference point (the original kept length).
    del total
    return kept, dropped


def _pop_windowed_edges(
    page_lines: List[List[str]],
    repeating_firsts: set,
    repeating_lasts: set,
    window: int = _PDF_EDGE_WINDOW,
) -> int:
    """Drop matching lines anywhere in the top-N / bottom-N window of each page.

    Returns the total number of lines removed across all pages so the
    caller can roll the count into structured ingestion notes for
    downstream visibility. Pages that end up empty after the pop are
    left as empty lists (the caller filters them out at the join step).
    Per-page work is delegated to :func:`_pop_one_page_window`.
    """
    stripped = 0
    for lines in page_lines:
        if not lines:
            continue
        kept, dropped = _pop_one_page_window(lines, repeating_firsts, repeating_lasts, window)
        stripped += dropped
        lines[:] = kept
    return stripped


def _strip_repeating_page_lines(pages: List[str]) -> Tuple[List[str], int]:
    """Strip page-edge lines that repeat across the document.

    Phase 15 Task 1 (audit §1.1):
        The pre-Phase-15 implementation collected only ``lines[0]`` and
        ``lines[-1]`` per pass and looped until no more recurrence met
        the 70 % cutoff. The mechanism was correct in principle (it did
        iterate) but the loop's exit condition killed it on a corpus
        with a *variable* outer line — a per-chapter section title on
        top, an incrementing page number on the bottom — because pass 1
        found no recurrence at the outermost slot and broke before
        reaching the deeper-row constant line that ought to have been
        stripped on pass 2. Widening the inspection to the top-N /
        bottom-N rows per page (default :data:`_PDF_EDGE_WINDOW` = 3)
        catches a line that recurs at any position within the window
        in a single pass.

    Why ``iterative-peel alone is not enough``:
        The same comment is duplicated inline above the loop body so a
        future implementer reading just the docstring cannot reintroduce
        the regression by collapsing the window back to 1. The bug is
        the loop's *exit condition*, not the outer-most-line check — a
        casual reviewer can easily mis-read "we already iterate" as
        "the algorithm is correct" and patch only the symptom.

    Returns:
        ``(cleaned_pages, lines_stripped)``. Empty pages (after stripping)
        are filtered out at the join step so they do not pollute the
        downstream paragraph chunker with phantom blank pages.

    The dedup is a no-op on documents with fewer than 3 pages — the
    statistical signal is too weak to distinguish "header" from
    "actual repeated paragraph".
    """
    if len(pages) < _PDF_REPEAT_MIN_PAGES:
        return pages, 0

    page_lines: List[List[str]] = [[ln.strip() for ln in p.splitlines() if ln.strip()] for p in pages]
    # ``math.ceil`` (not ``int``) so the 70 % rule actually fires at 70 %.
    # Example: 5 pages × 0.7 = 3.5 → ``int`` gave 3 (60 %, too lenient);
    # ``math.ceil`` gives 4 (80 %, ≥ 70 % as documented).
    cutoff = max(2, math.ceil(_PDF_REPEAT_THRESHOLD * len(pages)))
    total_stripped = 0

    # Pass 1 / pass 2 / … : window-based detection peels one layer per
    # pass. Iterating is still needed because once the outermost lines
    # are stripped the next-deepest row of the next pass might now sit
    # at offset 0 and meet the cutoff at the same window, but the
    # iteration is no longer *gated* on the strict-outermost recurrence —
    # so a variable-outer-line corpus is no longer locked out (audit §1.1).
    while True:
        repeating_firsts, repeating_lasts = _windowed_repeating_edge_lines(page_lines, cutoff)
        if not repeating_firsts and not repeating_lasts:
            break
        stripped = _pop_windowed_edges(page_lines, repeating_firsts, repeating_lasts)
        if stripped == 0:
            # Defensive break: no edge in the current window matched
            # despite the set being non-empty (e.g. every match was
            # already at depth > window). Avoid an infinite loop.
            break
        total_stripped += stripped

    cleaned = ["\n".join(lines) for lines in page_lines if lines]
    return cleaned, total_stripped


def strip_paragraph_packed_headers(text: str, *, threshold: float = _PDF_REPEAT_THRESHOLD) -> Tuple[str, int]:
    """Second-pass dedup against text already glued into paragraph blocks.

    Phase 15 Task 1, second part: after the page-level dedup runs and
    the paragraph chunker greedy-packs the result, a header line that
    survived (e.g. because it sat at row 4 on some pages and was
    therefore outside the window) can still re-appear at the start of
    several chunks. This helper looks at the first line of each
    ``\\n\\n``-separated block and strips any line that recurs in
    ``threshold`` fraction of blocks. Returns ``(cleaned_text, n_stripped)``.

    Operates on already-extracted text, so it never blocks on I/O and
    has no PDF-specific dependency. Skipping it would leave the
    Task 1 acceptance criterion fragile: the page-level dedup catches
    *most* runs but the audit recommends a second pass after chunker
    packing to mop up the residue.
    """
    blocks = text.split("\n\n")
    if len(blocks) < _PDF_REPEAT_MIN_PAGES:
        return text, 0
    first_lines: List[str] = []
    for block in blocks:
        block = block.strip()
        if not block:
            first_lines.append("")
            continue
        first_lines.append(block.splitlines()[0].strip())
    counts = Counter(ln for ln in first_lines if ln)
    cutoff = max(2, math.ceil(threshold * len([ln for ln in first_lines if ln])))
    repeating = {ln for ln, n in counts.items() if n >= cutoff}
    if not repeating:
        return text, 0

    stripped = 0
    rebuilt: List[str] = []
    for block in blocks:
        if not block.strip():
            rebuilt.append(block)
            continue
        block_lines = block.splitlines()
        # Drop any leading lines that match the repeating set; this
        # mops up the survivor headers that the page-level pass missed.
        idx = 0
        while idx < len(block_lines) and block_lines[idx].strip() in repeating:
            stripped += 1
            idx += 1
        rebuilt.append("\n".join(block_lines[idx:]))
    return "\n\n".join(rebuilt), stripped


def _post_extract_normalise(text: str, ctx: _ExtractContext, file_path: Path) -> str:
    """Apply Phase 15 post-extract transforms in a deterministic order.

    Order matters and is documented as a Phase 15 contract:

    1. **Glyph normalisation** (Task 3) — fix pypdf font-fallback
       artefacts before any other pass sees them. If the script-sanity
       check sees the corrupt glyphs first the warning would fire
       loudly even when the normaliser would have repaired them in the
       next step.
    2. **Script-sanity check** (Task 2) — run after normalisation so
       the warning fires only on residual corruption. The report is
       appended to ``ctx.script_sanity_reports`` whether or not it
       triggered, so the structured notes can record "we checked".
    3. **Operator strip-patterns** (Wave 2 Task 11) — applied before
       URL handling because a pattern might legitimately delete a
       URL-bearing line wholesale.
    4. **URL handling** (Wave 2 Task 14) — mask / strip per
       ``ctx.strip_urls_mode``.

    Front-matter heuristic (Task 13) and PDF-specific multi-column
    detection (Task 15) live inside :func:`_extract_pdf` itself
    because they need pypdf state and run before this post-extract
    pass.
    """
    if not text:
        return text

    # 1. Glyph normalisation. The function no-ops when profile == "none".
    # Phase 15 round-1 review (C-2): when the profile is active and
    # actually rewrites characters, surface the count at INFO so an
    # operator who didn't intend to apply a Turkish profile to a
    # Norwegian / Estonian corpus has a single log line to spot the
    # mismatch. The check compares character identity rather than full
    # text equality so a long-document round-trip stays cheap.
    if ctx.normalise_profile and ctx.normalise_profile != "none":
        # Round-3 review: use the exact-count helper from
        # ``_pypdf_normalise`` instead of the previous zip-diff heuristic,
        # which over- and under-counted on multi-char rule boundaries.
        substitutions = _count_normalise_substitutions(text, ctx.normalise_profile)
        normalised = _apply_normalise_profile(text, ctx.normalise_profile)
        if substitutions:
            logger.info(
                "Normalisation profile %r applied %d substitution(s) on '%s'. "
                "Pass --no-normalise-unicode or --normalise-profile none to disable.",
                ctx.normalise_profile,
                substitutions,
                file_path,
            )
        text = normalised

    # 2. Script-sanity check. Always invoked when a language hint is
    # configured; the helper itself silently no-ops on unknown hints so
    # the operator can pass an unsupported code without crashing.
    if ctx.language_hint:
        report = check_script_sanity(
            text,
            language_hint=ctx.language_hint,
            file_path=str(file_path),
            threshold=ctx.script_sanity_threshold,
            profile=ctx.normalise_profile,
        )
        ctx.script_sanity_reports.append(report)

    # 3. Operator strip-patterns.
    if ctx.strip_patterns:
        text, subs = _apply_strip_patterns(
            text,
            ctx.strip_patterns,
            timeout_s=ctx.strip_pattern_timeout,
        )
        ctx.strip_pattern_substitutions_total += subs

    # 4. URL handling.
    if ctx.strip_urls_mode in ("mask", "strip"):
        replacement = "[URL]" if ctx.strip_urls_mode == "mask" else ""
        new_text, n = _URL_PATTERN.subn(replacement, text)
        text = new_text
        ctx.urls_handled_total += n

    return text


def _try_pdf_decrypt(reader: Any, path: Path) -> None:
    """Attempt empty-password decrypt for owner-encrypted PDFs; raise ValueError otherwise.

    Empty-password decrypt covers the common "owner-encrypted but readable"
    case. When that fails OR the reader cannot honour the empty password,
    we surface a clear message pointing the operator to external tooling
    (qpdf / pdftk) — wiring a CLI password flag would put credentials in
    shell history, which is the wrong default.

    Decision is made on ``reader.decrypt()``'s **return value**, not on
    ``reader.is_encrypted``: pypdf ≥ 4.0 keeps ``is_encrypted`` reflecting
    the original-file state even after a successful decrypt, so checking
    that flag would reject correctly-decrypted owner-encrypted PDFs. The
    decrypt method returns a :class:`pypdf.PasswordType` enum
    (``IntEnum``) where ``NOT_DECRYPTED == 0`` (falsy) and
    ``USER_PASSWORD`` / ``OWNER_PASSWORD`` are truthy.
    """
    from pypdf.errors import DependencyError, FileNotDecryptedError

    try:
        result = reader.decrypt("")
    except (FileNotDecryptedError, NotImplementedError, DependencyError) as exc:
        raise ValueError(
            f"PDF '{path}' is encrypted. Decrypt it first (qpdf --decrypt / pdftk input_pw …) and re-run ingest."
        ) from exc
    if not result:
        # PasswordType.NOT_DECRYPTED == 0 → empty password didn't unlock it.
        raise ValueError(f"PDF '{path}' is encrypted with a non-empty password. Decrypt externally before ingest.")


def _read_pdf_pages(
    reader: Any,
    path: Path,
    *,
    page_range: Optional[Tuple[int, int]] = None,
) -> List[Tuple[int, str]]:
    """Extract per-page text; tolerate page-level failures with a loud warning.

    Page-level extraction failure is **not** propagated as a file-level
    error: a 500-page PDF with one malformed page should still produce
    499 good pages. We *do* upgrade the prior debug-log to a warning
    that names the path and page index so the operator sees in CI logs
    exactly which page failed and can spot-check it. The post-loop
    "no extractable text" warning still catches the all-pages-failed
    case downstream in :func:`_extract_pdf`.

    ``page_range`` (1-indexed inclusive) lets the caller restrict
    extraction to a subset of pages. The validator in
    :func:`_validate_page_range` enforces the bounds before we reach
    this helper; here we only slice. Returns ``(page_index, text)`` pairs
    so downstream callers (front-matter heuristic, multi-column warning)
    can refer to the original page numbers in their reporting even after
    the slice.
    """
    total = len(reader.pages)
    if page_range is not None:
        start_1idx, end_1idx = page_range
        # Clamp end to total so a range of 5-9999 on a 200-page doc
        # silently shrinks to 5-200; the explicit out-of-range error
        # comes from _validate_page_range, which runs at CLI dispatch.
        end_1idx = min(end_1idx, total)
        indices = range(start_1idx - 1, end_1idx)
    else:
        indices = range(total)

    pages: List[Tuple[int, str]] = []
    for idx in indices:
        page = reader.pages[idx]
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 — best-effort: pypdf's per-page extraction surface is wide (KeyError on malformed object refs, AssertionError on broken cross-ref tables, UnicodeDecodeError on font encodings, plus its own internal errors); per-page soft-fail keeps the run going so a single bad page cannot abort a multi-thousand-document corpus ingest.  # NOSONAR
            logger.warning(
                "PDF page extraction failed (file=%s, page_index=%d): %s — page skipped, run continues.",
                path,
                idx,
                exc,
            )
            text = ""
        if text.strip():
            pages.append((idx, text))
    return pages


def _validate_page_range(page_range: Tuple[int, int], total_pages: int, path: Path) -> Tuple[int, int]:
    """Validate ``page_range`` against ``total_pages``.

    Raises :class:`IngestParameterError` (a :class:`ValueError` subclass)
    with an operator-facing message on failure. The CLI dispatcher
    translates this to ``EXIT_CONFIG_ERROR`` per the existing per-
    subcommand pattern. Library callers see the raw exception, in line
    with the documented contract.

    Why :class:`IngestParameterError` and not plain :class:`ValueError`:
    the per-file extractor's soft-fail catch in
    :func:`_extract_text_for_ingest` swallows generic exceptions to keep
    a multi-thousand-file corpus running when one bad page or one
    locked file is in the way. An operator-supplied parameter mistake
    is a different class of failure — it should abort the run loudly
    rather than fold into the log. The custom subclass lets the catch
    re-raise this one specifically.
    """
    start, end = page_range
    if start < 1:
        raise IngestParameterError(f"--page-range start ({start}) must be >= 1 for '{path}' (1-indexed).")
    if start > end:
        raise IngestParameterError(f"--page-range start ({start}) must be <= end ({end}) for '{path}'.")
    if start > total_pages:
        raise IngestParameterError(f"--page-range start ({start}) exceeds PDF page count ({total_pages}) for '{path}'.")
    return (start, min(end, total_pages))


# ---------------------------------------------------------------------------
# Phase 15 Task 13 — front-matter heuristic for PDFs
# ---------------------------------------------------------------------------

_FRONTMATTER_PROBE_PAGES: int = 12
"""Number of leading / trailing pages probed by the front-matter heuristic.

The 12-page envelope matches the audit's empirical observation: ToC /
front-matter material in a 200-page corpus rarely exceeds the first
dozen pages, and the symmetric trailing pass catches index / glossary
back-matter under the same envelope. Configurable in future phases if
operator feedback shifts the modal value.
"""

_FRONTMATTER_ALPHA_RATIO_MAX: float = 0.45
_FRONTMATTER_LEADER_RATIO_MIN: float = 0.10
_FRONTMATTER_PAGE_NUM_PATTERN = re.compile(r"\n\d{1,3}\n")
# Round-3 review: ToC pages commonly use **dotted** leaders (`.....`)
# OR underscore leaders (`____`). A bare-underscore-only count missed
# the dot-leader case on real-world publications, leaving genuine ToC
# pages slipping past the heuristic. Combine both leader styles into a
# single regex; runs of length ≥ 3 reliably distinguish leaders from
# legitimate punctuation.
_FRONTMATTER_LEADER_RUN_PATTERN = re.compile(r"[._]{3,}")


def _is_frontmatter_page(text: str) -> bool:
    """Apply the audit's three-signal heuristic: low alpha + high leader ratio + many page-numbers.

    A page is treated as "front-matter candidate" iff **all three**
    conditions hold:

    * Alphabetic-character ratio < 0.45.
    * Leader-character ratio > 0.10. Round-3 fix: the leader count
      now covers BOTH underscore runs (``_____``) AND dot runs
      (``.....``) of length ≥ 3, denominated over the non-whitespace
      character count (parity with ``alpha_ratio``). Pre-round-3 the
      check only counted single ``_`` characters against the raw text
      length, which missed the dotted-leader ToC pages that the
      docstring already promised to handle.
    * At least 5 inline-page-number matches (``\\n<1-3 digits>\\n``).

    Three independent signals keep the false-positive rate low — body
    text rarely meets even two of them simultaneously.
    """
    if not text:
        return False
    non_ws = [c for c in text if not c.isspace()]
    if not non_ws:
        return False
    alpha_ratio = sum(1 for c in non_ws if c.isalpha()) / len(non_ws)
    leader_chars = sum(len(m) for m in _FRONTMATTER_LEADER_RUN_PATTERN.findall(text))
    leader_ratio = leader_chars / len(non_ws)
    page_num_hits = len(_FRONTMATTER_PAGE_NUM_PATTERN.findall(text))
    return (
        alpha_ratio < _FRONTMATTER_ALPHA_RATIO_MAX
        and leader_ratio > _FRONTMATTER_LEADER_RATIO_MIN
        and page_num_hits >= 5
    )


def _drop_frontmatter_pages(
    pages: List[Tuple[int, str]],
    probe: int = _FRONTMATTER_PROBE_PAGES,
) -> Tuple[List[Tuple[int, str]], List[int]]:
    """Drop leading + trailing front-matter pages; return the remainder + dropped indices.

    Symmetric pass: the first ``probe`` pages are dropped from the
    head while they keep matching the heuristic; the last ``probe``
    pages similarly from the tail. Stops at the first non-matching
    page in each direction so a real body page can never be dropped
    mid-document.
    """
    if not pages:
        return pages, []
    dropped: List[int] = []
    head_keep_idx = 0
    for pg_idx, (orig_idx, text) in enumerate(pages[:probe]):
        if _is_frontmatter_page(text):
            dropped.append(orig_idx)
            head_keep_idx = pg_idx + 1
        else:
            break
    tail_keep_idx = len(pages)
    for offset in range(min(probe, len(pages) - head_keep_idx)):
        pg = pages[-(offset + 1)]
        orig_idx, text = pg
        if _is_frontmatter_page(text):
            dropped.append(orig_idx)
            tail_keep_idx = len(pages) - (offset + 1)
        else:
            break
    remaining = pages[head_keep_idx:tail_keep_idx]
    return remaining, sorted(set(dropped))


# ---------------------------------------------------------------------------
# Phase 15 Task 15 — multi-column PDF detection (warning only)
# ---------------------------------------------------------------------------

_MULTI_COLUMN_SEPARATION_PCT: float = 0.30
"""Fraction-of-page-width gap between text clusters that flags 2-column.

Round-2 review (nit on ingestion.py:700-758) proposed lowering this from
0.30 → 0.08 to match real-world 5-8 %-of-page-width gutters. After
investigation we **kept the 0.30 threshold** because the current
detector's cluster split is computed against the min / max of all
extracted text-matrix x-offsets, NOT against per-line start positions.
pypdf's ``visitor_text`` callback fires per-glyph with the running text
matrix, so a single-line "Body content here." emits x-positions spanning
the line's horizontal extent (e.g. 72 → 250 on a US-letter page) — the
min / max gap on a single-column page can already exceed 25 %, well
above the proposed 8 % threshold. A genuine 2-column page has its
right-column glyphs starting near the page midline (≈ 300 / 612 ≈ 50 %)
so 0.30 catches it without false-positiving on single-column body text.

A histogram-based bimodal-mode detector would let us safely drop the
threshold below 0.10; that refactor is out of scope for the Phase 15
round-2 fix and is tracked as a Wave 3 follow-up (see
`docs/roadmap/phase-15-ingestion-reliability.md` Wave 3 — multi-column
layout extraction).
"""


def _select_sample_pages(reader: Any, page_range: Optional[Tuple[int, int]]) -> List[Any]:
    """Slice the first three text-bearing pages for the multi-column probe.

    Honours ``--page-range`` when set so the warning reflects the real
    extraction surface (round-2 review nit). Returns an empty list on
    any slicing failure — the probe degrades to silence rather than
    blocking the run.
    """
    try:
        if page_range is not None:
            start_1idx, end_1idx = page_range
            return reader.pages[start_1idx - 1 : end_1idx][:3]
        return reader.pages[:3]
    except Exception:  # noqa: BLE001 — pypdf wrapper objects may raise on slice on torn PDFs.
        return []


def _collect_multi_column_samples(sample_pages: List[Any]) -> Tuple[List[float], Optional[float]]:
    """Drive pypdf's ``visitor_text`` callback across ``sample_pages``.

    Returns ``(x_positions, page_width)``. The page width comes from the
    first page that exposes a ``mediabox``. Sampling stops once 40
    positions have been collected — enough to call multi-column with
    confidence without walking a 500-page document.
    """
    x_positions: List[float] = []
    page_width: Optional[float] = None

    def visitor(text: str, _cm: Any, tm: Any, _fontdict: Any, _fontsize: Any) -> None:
        # ``tm`` is a 6-tuple ``(a, b, c, d, e, f)`` representing the
        # text-matrix; the horizontal offset is element 4. Pypdf calls
        # this with the documented kwargs (``font_dict`` / ``font_size``
        # under the historical CamelCase names) — the snake_case
        # underscored names here keep Sonar python:S117 happy without
        # changing the call contract.
        if text and tm is not None:
            try:
                x_positions.append(float(tm[4]))
            except (TypeError, IndexError, ValueError):
                return

    for page in sample_pages:
        try:
            page.extract_text(visitor_text=visitor)
            if page_width is None and getattr(page, "mediabox", None) is not None:
                page_width = float(page.mediabox.width)
        except Exception:  # noqa: BLE001 — sampling is best-effort; per-page failures from torn PDFs / unusable font tables should not block the multi-column probe.  # nosec B112
            continue
        if len(x_positions) >= 40:
            break
    return x_positions, page_width


def _is_two_cluster_distribution(x_positions: List[float], page_width: float) -> bool:
    """Return ``True`` iff the x-positions split into two clusters past the threshold."""
    if not x_positions or page_width <= 0:
        return False
    sorted_xs = sorted(x_positions)
    midpoint = (sorted_xs[0] + sorted_xs[-1]) / 2.0
    left_cluster_max = max((x for x in sorted_xs if x <= midpoint), default=midpoint)
    right_cluster_min = min((x for x in sorted_xs if x > midpoint), default=midpoint)
    gap = right_cluster_min - left_cluster_max
    return gap > _MULTI_COLUMN_SEPARATION_PCT * page_width


def _maybe_warn_multi_column(reader: Any, path: Path, *, page_range: Optional[Tuple[int, int]] = None) -> bool:
    """Sample the first text-bearing pages for a two-cluster x-distribution.

    Walks pypdf's text fragments via the ``visitor_text`` callback,
    accumulates per-fragment x-coordinates, and reports a two-cluster
    distribution as a single ``WARNING``. No fix attempt — the
    operator is expected to switch to ``--strategy sliding`` with a
    larger chunk-size or to pre-process the PDF with a layout-aware
    tool. Returns ``True`` iff a warning was emitted (test seam).

    Round-2 review (nit on ingestion.py:803): when ``--page-range`` is
    set the warning previously sampled pages 0-2 even if those pages
    were going to be dropped from the actual extraction; we now respect
    the same slice the extractor will use so the warning reflects the
    real run, not the file's first three pages.
    """
    sample_pages = _select_sample_pages(reader, page_range)
    if not sample_pages:
        return False
    x_positions, page_width = _collect_multi_column_samples(sample_pages)
    if page_width is None:
        return False
    if not _is_two_cluster_distribution(x_positions, page_width):
        return False
    logger.warning(
        "Detected 2-column layout in '%s' — reading order may be scrambled. "
        "Consider --strategy sliding with a larger --chunk-size, or pre-process "
        "the PDF with a layout-aware tool (camelot-py / pdfplumber).",
        path,
    )
    return True


def _extract_pdf(path: Path, *, ctx: Optional[_ExtractContext] = None) -> str:
    """Extract PDF text with Phase 15's hardened pipeline.

    Phase 15 additions vs pre-15:

    * ``page_range`` filtering (Wave 2 Task 12) before per-page extraction.
    * Multi-column detection warning (Wave 2 Task 15) sampled on the
      first three pages of the slice.
    * Front-matter heuristic drop (Wave 2 Task 13) on the first 12 / last
      12 pages of the slice, opt-out via ``ctx.keep_frontmatter``.
    * Glyph-normalisation + script-sanity + strip-pattern + strip-URLs
      are applied by :func:`_post_extract_normalise` after this function
      returns, so this body deliberately stays PDF-specific.
    """
    if ctx is None:
        ctx = _ExtractContext()
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:  # pragma: no cover — covered by extras
        # Narrow on ModuleNotFoundError + name match so we only convert
        # genuine "extra not installed" failures.  A corrupt install or
        # circular-import in pypdf raises a plain ImportError (or
        # ModuleNotFoundError with a different ``name``); re-raising those
        # preserves the original traceback for debugging.
        if exc.name != "pypdf":
            raise
        raise OptionalDependencyError(
            "PDF ingestion requires the 'ingestion' extra. Install with: pip install 'forgelm[ingestion]'"
        ) from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001 — best-effort: PdfReader open surfaces OSError (file/permission), pypdf-internal errors (bad header, unsupported version, malformed xref), and on rare adversarial inputs InfiniteLoopError; converting all to a typed ValueError keeps the per-file error path uniform with DOCX/EPUB extractors.  # NOSONAR
        raise ValueError(f"Could not open PDF '{path}': {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        _try_pdf_decrypt(reader, path)

    page_range = ctx.page_range
    if page_range is not None:
        page_range = _validate_page_range(page_range, len(reader.pages), path)

    _maybe_warn_multi_column(reader, path, page_range=page_range)

    indexed_pages = _read_pdf_pages(reader, path, page_range=page_range)
    if not indexed_pages:
        logger.warning(
            "No extractable text in '%s'. Likely a scanned PDF without a text layer; "
            "run OCR (Tesseract / AWS Textract) before ingest.",
            path,
        )
        return ""

    if not ctx.keep_frontmatter:
        indexed_pages, dropped = _drop_frontmatter_pages(indexed_pages)
        if dropped:
            ctx.frontmatter_dropped_pages.extend(dropped)
            logger.warning(
                "Dropped %d front-matter / back-matter page(s) from '%s' "
                "(indices=%s). Use --keep-frontmatter to retain them.",
                len(dropped),
                path,
                dropped,
            )

    page_texts = [text for _, text in indexed_pages]
    cleaned_pages, stripped = _strip_repeating_page_lines(page_texts)
    if ctx.dedup_state is not None and stripped:
        # Roll into the run-level structured notes so the operator can
        # tell post-hoc that header/footer dedup was actually doing work.
        ctx.dedup_state["lines_stripped"] = ctx.dedup_state.get("lines_stripped", 0) + stripped

    joined = "\n\n".join(cleaned_pages)
    joined, secondary_stripped = strip_paragraph_packed_headers(joined)
    if ctx.dedup_state is not None and secondary_stripped:
        ctx.dedup_state["paragraph_packed_stripped"] = (
            ctx.dedup_state.get("paragraph_packed_stripped", 0) + secondary_stripped
        )
    return joined


def _escape_md_cell(text: Optional[str]) -> str:
    """Escape ``|`` and ``\\`` so cell text is safe inside a markdown row.

    Newlines collapse to spaces — a multi-line cell can't be expressed
    inside one markdown table row. Empty / ``None`` returns ``""``.
    """
    if not text:
        return ""
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("\n", " ").replace("\r", " ").strip()


def _docx_table_to_markdown(table: Any) -> str:
    """Phase 12: render a python-docx ``Table`` as a markdown-table block.

    The previous behaviour was to flatten every row to ``" | "`` and lose
    the header/separator distinction. The markdown form keeps the
    header-row signal intact, which matters for SFT use cases (tabular
    Q&A, financial assistant, code-with-data) where a "first row is the
    header" cue is meaningful for the model.

    Empty rows / empty trailing cells are stripped; rows with mismatched
    column counts (rare; usually a merged-cell artefact) are padded
    with empty cells so the markdown stays well-formed. The first
    non-empty row is treated as the header (no heuristic — that's what
    DOCX authors mean when they put a row in the table's first slot).

    Cell text containing markdown structural characters (``|`` and ``\\``)
    is escaped per CommonMark — otherwise a cell value like ``a|b`` would
    be parsed as two extra columns by downstream tokenisers / renderers.
    Newlines inside cells are collapsed to spaces (a multi-line cell
    can't be expressed in a single markdown table row).
    """
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [_escape_md_cell(cell.text) for cell in row.cells]
        # Trim purely-empty rows; a row of all blanks is usually a layout artefact.
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    header = rows[0]
    body = rows[1:]
    header_line = "| " + " | ".join(header) + " |"
    separator_line = "|" + "|".join("---" for _ in range(width)) + "|"
    body_lines = ["| " + " | ".join(row) + " |" for row in body]

    return "\n".join([header_line, separator_line, *body_lines])


def _iter_docx_blocks(doc: Any) -> Iterable[Any]:
    """Yield block-level elements (paragraphs + tables) in document order.

    ``python-docx``'s ``doc.paragraphs`` and ``doc.tables`` collections lose
    the relative ordering of headings → tables → paragraphs that authors
    rely on. Walking the underlying ``<w:body>`` XML in document order
    preserves layout — critical for tabular Q&A and policy manuals where
    the table is supposed to appear *after* the paragraph that introduces
    it, not at the end of the file.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _add_part_lines_to_boilerplate(part: Any, boilerplate: set) -> None:
    """Drain non-empty stripped lines from a header/footer part into ``boilerplate``."""
    for paragraph in getattr(part, "paragraphs", []):
        text = getattr(paragraph, "text", "") or ""
        for ln in text.splitlines():
            cleaned = ln.strip()
            if cleaned:
                boilerplate.add(cleaned)


def _collect_docx_header_footer_lines(doc: Any) -> set:
    """Phase 15 Task 6 — return every line declared in a section's header / footer.

    Word documents declare headers / footers explicitly under each
    section's ``<w:hdr>`` / ``<w:ftr>`` parts. python-docx exposes them
    as ``doc.sections[i].header.paragraphs`` and ``.footer.paragraphs``.
    We collect their non-empty stripped lines into a single ``set`` so a
    multi-section document with the same boilerplate on each section
    only contributes one entry per unique line. Body extraction
    afterwards subtracts these lines, eliminating the
    "header bleeds into body chunk" failure mode the audit (§3.2)
    documented. Per-section drain logic is delegated to
    :func:`_add_part_lines_to_boilerplate` to keep this function below
    Sonar S3776's cognitive-complexity ceiling.

    Defensive: a malformed section (no header element) raises
    ``AttributeError`` from python-docx; we swallow it per-section so a
    single bad section cannot kill extraction of the rest.
    """
    boilerplate: set = set()
    for section in getattr(doc, "sections", []):
        for source_name in ("header", "footer"):
            try:
                part = getattr(section, source_name, None)
            except AttributeError:
                continue
            if part is None:
                continue
            _add_part_lines_to_boilerplate(part, boilerplate)
    return boilerplate


_DOCX_BOILERPLATE_MAX_LEN: int = 80
"""Max stripped-line length eligible for boilerplate subtraction (round-2 S-4).

A body paragraph beginning with ``Chapter 1`` is overwhelmingly longer
than its standalone-title-as-header counterpart, so the length floor
protects body text from being mistakenly dropped along with the
running-header line of the same text.
"""


def _strip_docx_boilerplate(text: str, boilerplate: set) -> Optional[str]:
    """Return ``text`` minus header/footer boilerplate lines, or ``None`` if empty.

    Splits the paragraph into lines, drops any line that is **short
    enough** to plausibly be a header (≤ 80 chars) AND exactly matches
    a header / footer entry. Returns the cleaned text on success, or
    ``None`` when nothing survives (caller skips the paragraph).
    """
    kept = [
        ln
        for ln in text.splitlines()
        if not (len(ln.strip()) <= _DOCX_BOILERPLATE_MAX_LEN and ln.strip() in boilerplate)
    ]
    cleaned = "\n".join(kept).strip()
    return cleaned or None


def _render_docx_block(element: Any, boilerplate: set, table_cls: Any) -> Optional[str]:
    """Render one ``_iter_docx_blocks`` element as text, or ``None`` to skip."""
    if isinstance(element, table_cls):
        rendered = _docx_table_to_markdown(element)
        return rendered or None
    text = getattr(element, "text", "")
    if not text or not text.strip():
        return None
    if not boilerplate:
        return text
    return _strip_docx_boilerplate(text, boilerplate)


def _extract_docx(path: Path, *, ctx: Optional[_ExtractContext] = None) -> str:
    """Extract DOCX text with explicit header / footer subtraction (Phase 15 Task 6).

    Round-2 review (S-4 + Sonar S3776): per-block rendering is delegated
    to :func:`_render_docx_block` so this function stays a thin
    open-doc → iterate-blocks → join loop below the cognitive-complexity
    ceiling. The pre-review code subtracted EVERY paragraph line whose
    stripped form appeared in the header / footer set, which dropped
    legitimate body paragraphs matching a header verbatim
    (``Chapter 1`` body vs ``Chapter 1`` header). The new helper now
    short-circuits on lines longer than :data:`_DOCX_BOILERPLATE_MAX_LEN`.
    """
    if ctx is None:
        ctx = _ExtractContext()
    try:
        from docx import Document
        from docx.table import Table
    except ModuleNotFoundError as exc:  # pragma: no cover
        # See PDF block above for the narrowing rationale.  ``python-docx``
        # imports as ``docx``; ``docx.table`` is in the same package so a
        # missing-extra failure surfaces as exc.name == "docx".
        if exc.name not in ("docx", "docx.table"):
            raise
        raise OptionalDependencyError(
            "DOCX ingestion requires the 'ingestion' extra. Install with: pip install 'forgelm[ingestion]'"
        ) from exc

    doc = Document(str(path))
    boilerplate = _collect_docx_header_footer_lines(doc)
    blocks: List[str] = []
    for element in _iter_docx_blocks(doc):
        rendered = _render_docx_block(element, boilerplate, Table)
        if rendered:
            blocks.append(rendered)
    return "\n\n".join(blocks)


_EPUB_NAME_TOKEN_SPLIT = re.compile(r"[/\\.\-_ ]+")


def _epub_item_matches_skip(item_name: str, item_type: str, skip_list: Tuple[str, ...]) -> bool:
    """Return ``True`` iff the item's file name or epub:type matches the skip-list.

    Phase 15 round-1 review (C-1): the pre-review implementation used a
    plain substring match, which silently skipped legitimate chapters
    whose filenames *contained* a skip token — ``recovery.xhtml`` got
    dropped because ``cover`` is a substring of ``recovery``;
    ``navy.xhtml`` and ``discovery_chapter.xhtml`` had the same fate.
    A novel with chapters named "Recovery from War" or "The Royal Navy"
    would silently lose entire bodies of text with no operator-facing
    warning — the same silent-failure shape Phase 15 was launched to
    fix.

    The fix is a **whole-token** match. We split the file name on
    common path / extension / separator characters and compare each
    resulting token against the skip-list. This catches the canonical
    cases (``cover.xhtml``, ``nav.xhtml``, ``copyright.xhtml``,
    ``cover-page.xhtml``, ``oebps/cover.xhtml``) while keeping
    ``recovery.xhtml`` / ``navy.xhtml`` / ``discovery_chapter.xhtml``
    safe.

    Exact-token match on the ``epub:type`` value is unchanged.
    """
    name_lc = item_name.lower()
    tokens = {token for token in _EPUB_NAME_TOKEN_SPLIT.split(name_lc) if token}
    if tokens & set(skip_list):
        return True
    type_lc = (item_type or "").lower()
    return bool(type_lc) and type_lc in skip_list


def _resolve_spine_item(book: Any, spine_entry: Any, item_document_type: int) -> Optional[Tuple[Any, str, str]]:
    """Look up the spine entry's manifest item; return ``(item, name, type_str)`` or ``None``.

    ``None`` is returned for missing items or non-document types so the
    caller can ``continue``. The ``type_str`` is the space-joined EPUB-3
    manifest properties (e.g. ``"nav cover-image"``) — empty for items
    without explicit properties.
    """
    item_id = spine_entry[0] if isinstance(spine_entry, (tuple, list)) else spine_entry
    item = book.get_item_with_id(item_id)
    if item is None or item.get_type() != item_document_type:
        return None
    item_name = getattr(item, "file_name", "") or getattr(item, "name", "") or ""
    # Round-2 nit: ebooklib exposes EPUB-3 manifest properties (e.g.
    # ``properties=["nav", "cover-image"]``) on ``item.properties``;
    # the pre-round-2 code probed ``media_overlay`` which is always
    # empty for nav / cover items. Join with spaces so any property
    # token can match the skip-list via the whole-token splitter
    # introduced for C-1.
    item_properties = getattr(item, "properties", None) or []
    item_type = " ".join(str(p) for p in item_properties) if item_properties else ""
    return item, item_name, item_type


def _emit_epub_skip_warning(path: Path, skipped_items: List[str]) -> None:
    """Emit the Phase 15 round-1 C-1 WARNING naming skipped frontmatter items."""
    if not skipped_items:
        return
    preview = ", ".join(skipped_items[:6]) + (", …" if len(skipped_items) > 6 else "")
    logger.warning(
        "EPUB '%s': skipped %d frontmatter / navigation item(s): %s. "
        "Pass --epub-no-skip-frontmatter to keep these in the JSONL.",
        path,
        len(skipped_items),
        preview,
    )


def _extract_epub(path: Path, *, ctx: Optional[_ExtractContext] = None) -> str:
    """Extract EPUB text in reading order; skip nav / cover / copyright (Phase 15 Task 7).

    Spine resolution and the skip-WARNING surface are extracted into
    ``_resolve_spine_item`` and ``_emit_epub_skip_warning`` to keep
    this function below the Sonar S3776 cognitive-complexity ceiling.
    """
    if ctx is None:
        ctx = _ExtractContext()
    try:
        from bs4 import BeautifulSoup
        from ebooklib import ITEM_DOCUMENT, epub
    except ModuleNotFoundError as exc:  # pragma: no cover
        # See PDF block above for the narrowing rationale.  EPUB ingestion
        # depends on ``beautifulsoup4`` (imported as ``bs4``) and
        # ``ebooklib``; ``ebooklib.epub`` is a submodule so a missing extra
        # surfaces as one of these three names.
        if exc.name not in ("bs4", "ebooklib", "ebooklib.epub"):
            raise
        raise OptionalDependencyError(
            "EPUB ingestion requires the 'ingestion' extra. Install with: pip install 'forgelm[ingestion]'"
        ) from exc

    # ebooklib >= 0.18 emits deprecation warnings unless `options=` is passed
    # explicitly. ignore_ncx=True silences the noisy NCX (table of contents)
    # warning that adds nothing for ingestion. ignore_missing_css avoids
    # CSS-resolution warnings on EPUBs that ship broken stylesheets.
    book = epub.read_epub(
        str(path),
        options={"ignore_ncx": True, "ignore_missing_css": True},
    )

    chunks: List[str] = []
    skipped_items: List[str] = []
    for spine_entry in getattr(book, "spine", []):
        resolved = _resolve_spine_item(book, spine_entry, ITEM_DOCUMENT)
        if resolved is None:
            continue
        item, item_name, item_type = resolved
        if ctx.epub_skip_frontmatter and _epub_item_matches_skip(item_name, item_type, ctx.epub_skip_items):
            skipped_items.append(item_name)
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text(separator="\n").strip()
        if text:
            chunks.append(text)
    _emit_epub_skip_warning(path, skipped_items)
    return "\n\n".join(chunks)


def _read_text_with_bom_strip(path: Path) -> str:
    """Phase 15 Task 8 — read a TXT / MD file and strip a leading UTF-8 BOM.

    Why a dedicated helper: the previous behaviour passed the BOM through
    as a literal character into chunk 0; downstream tokenisers rarely
    handle U+FEFF correctly. We use ``encoding="utf-8-sig"`` for the
    initial decode attempt — it strips the BOM transparently — and fall
    back to ``encoding="utf-8"`` with ``errors="replace"`` when the
    file is genuinely not UTF-8 (so the existing binary-contamination
    warning still fires).

    Round-2 review (S-2): the fallback path used to read with
    ``encoding="utf-8"`` rather than ``"utf-8-sig"``, leaking a literal
    ``\\ufeff`` into chunk 0 whenever a file mixed a BOM with downstream
    non-UTF-8 bytes. We now use ``"utf-8-sig"`` on both paths and
    additionally strip an explicit leading ``\\ufeff`` so any encoding-
    detection edge case is caught belt-and-braces.
    """
    try:
        raw = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="utf-8-sig", errors="replace")
    if raw.startswith("﻿"):
        raw = raw[1:]
    return raw


def _warn_if_binary_contamination(raw: str, path: Path) -> None:
    """Existing binary-content warning, factored out for reuse."""
    if raw:
        replacement_count = raw.count("�")
        if replacement_count / max(len(raw), 1) > 0.01:
            logger.warning(
                "'%s' contains %d Unicode replacement chars (%.1f%% of file). "
                "Likely binary content masquerading as text — verify the file is actually UTF-8.",
                path,
                replacement_count,
                replacement_count * 100 / len(raw),
            )


def _extract_text(path: Path, *, ctx: Optional[_ExtractContext] = None) -> str:
    """Extract plain TXT with UTF-8 BOM stripping (Phase 15 Task 8).

    The ``ctx`` parameter is accepted for signature parity with the
    other format extractors (the dispatcher calls every extractor with
    ``ctx=ctx``) and intentionally unused inside this body — TXT has
    no PDF / DOCX / EPUB-style options that need threading through.
    """
    del ctx  # signature parity only; Sonar S1854.
    raw = _read_text_with_bom_strip(path)
    _warn_if_binary_contamination(raw, path)
    return raw


def _extract_markdown(path: Path, *, ctx: Optional[_ExtractContext] = None) -> str:
    """Extract Markdown with UTF-8 BOM + YAML-frontmatter stripping (Phase 15 Task 8).

    YAML frontmatter is the ``---\\n…\\n---\\n`` block at the start of
    the file used by Jekyll / Hugo / MkDocs / Pelican. Without
    stripping it, chunk 0 of every imported MD file leads with
    ``title:`` / ``date:`` / ``author:`` rows that are pure metadata
    noise for SFT. Opt-out via ``ctx.keep_md_frontmatter`` for
    workflows that genuinely want to train on the metadata.
    """
    if ctx is None:
        ctx = _ExtractContext()
    raw = _read_text_with_bom_strip(path)
    _warn_if_binary_contamination(raw, path)
    if not ctx.keep_md_frontmatter:
        match = _YAML_FRONTMATTER_PATTERN.match(raw)
        if match is not None:
            raw = raw[match.end() :]
    return raw


_EXTRACTORS: dict = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".epub": _extract_epub,
    ".txt": _extract_text,
    ".md": _extract_markdown,
}


# ---------------------------------------------------------------------------
# Chunking strategies — each yields chunk strings from raw text
# ---------------------------------------------------------------------------


def _chunk_sliding(text: str, chunk_size: int, overlap: int) -> Iterable[str]:
    """Fixed character window with overlap. Coarse but predictable."""
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap must be in [0, chunk_size)")
    # Reject pathological overlap up front: ratios above 0.5 explode chunk
    # count (overlap=199 with chunk_size=200 yields ~one chunk per character).
    # The safety net is preventative — without it a typo can produce a
    # multi-million-line JSONL silently.
    if overlap > chunk_size // 2:
        raise ValueError(
            f"overlap ({overlap}) must be at most chunk_size // 2 ({chunk_size // 2}) to avoid quadratic chunk count. "
            "Reduce --overlap or increase --chunk-size."
        )
    if not text:
        return
    step = chunk_size - overlap
    for start in range(0, len(text), step):
        chunk = text[start : start + chunk_size]
        if chunk.strip():
            yield chunk
        # Stop once the current window already covers end-of-text. Without
        # this guard, large `--overlap` produces runt trailing chunks that
        # are pure prefixes of the prior chunk — they pollute the JSONL
        # with semantically-empty rows and skew the audit's near-duplicate
        # / length stats.
        if start + chunk_size >= len(text):
            return


def _chunk_paragraph(text: str, max_chunk_size: int) -> Iterable[str]:
    """Greedy paragraph packer — never splits a paragraph mid-sentence.

    Paragraphs longer than ``max_chunk_size`` are emitted on their own
    (caller's chunk_size becomes a soft cap). Keeps sentence boundaries
    intact so SFT examples don't start mid-thought.
    """
    if max_chunk_size <= 0:
        raise ValueError(_CHUNK_SIZE_POSITIVE_MSG)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        return

    current: List[str] = []
    current_len = 0
    for para in paragraphs:
        para_len = len(para)
        if current_len + para_len + 2 <= max_chunk_size or not current:
            current.append(para)
            current_len += para_len + 2
        else:
            yield "\n\n".join(current)
            current = [para]
            current_len = para_len
    if current:
        yield "\n\n".join(current)


def _chunk_semantic(text: str, chunk_size: int) -> Iterable[str]:
    # Embedding-based semantic chunking was deliberately deferred past
    # Phase 12 (see docs/roadmap/completed-phases.md — the
    # "deferred to Phase 13+" entry on embedding-based semantic dedup)
    # because a runtime embedding-model dependency conflicts with the
    # air-gapped Annex IV reproducibility guarantee. The closure plan
    # tracks the eventual ship vehicle as C-52 with a new optional
    # [chunking-semantic] extra. The placeholder issue tracking follow-up
    # work is recorded in docs/roadmap/risks-and-decisions.md under
    # F-PR29-A6-05-issue-link.
    raise NotImplementedError(
        "Semantic chunking requires an embedding model and is deferred past "
        "Phase 12 (see docs/roadmap/completed-phases.md — "
        "embedding-based semantic dedup deferred to Phase 13+ for Annex IV "
        "reproducibility); tracked under F-PR29-A6-05-issue-link in "
        "docs/roadmap/risks-and-decisions.md. Use 'sliding' or 'paragraph' "
        "for now."
    )


# ---------------------------------------------------------------------------
# Phase 12: markdown-aware splitter (heading + code-block boundary preserving)
# ---------------------------------------------------------------------------


# CommonMark allows 0-3 leading spaces before an ATX heading marker; 4+
# spaces would make the line an indented code block instead. The body
# capture is anchored on **non-whitespace** at both ends — without those
# anchors, the lazy ``(.+?)`` and greedy ``[ \t]*$`` tails compete for
# trailing whitespace, giving the engine O(n²) splits to try on a line
# like ``# x \t \t … \tx`` (~100 ms at n=2000 in CPython, blocking the
# ingest pipeline). The negated ``[^\n]*`` middle stays linear because
# it can't overlap with the anchoring ``\S`` on either side.
_MARKDOWN_HEADING_PATTERN = re.compile(
    r"^ {0,3}(#{1,6})[ \t]+(\S(?:[^\n]*\S)?)[ \t]*$",
    re.MULTILINE,
)
# Fence detection is a non-regex parser (see ``_parse_md_fence``) — the
# previous regex ``^ {0,3}(?P<fence>`{3,}|~{3,})(?P<rest>[^\n]*)$``
# satisfied SonarCloud python:S5852's "two unbounded greedy quantifiers
# in sequence" rule the same way our other markdown helpers do (state
# machines beat regexes for multi-line / per-line markdown parsing —
# see docs/standards/regex.md rule 6).


def _parse_md_fence(line: str) -> Optional[Tuple[str, int, str]]:
    """Detect a CommonMark fence line; return ``(fence_char, run_len, rest)`` or ``None``.

    CommonMark §4.5: a fence is 3+ identical backticks (or tildes),
    optionally indented up to 3 spaces. The portion after the fence run
    is the **info string** (allowed for openers, must be empty for
    closers — that policy is enforced by callers, not here).

    Implemented as a non-regex parser so the cost is provably O(n) per
    line. The previous regex
    ``^ {0,3}(?P<fence>\\`{3,}|~{3,})(?P<rest>[^\\n]*)$`` had two
    unbounded greedy quantifiers in sequence; SonarCloud python:S5852
    flags that shape as polynomial-runtime risk even when CPython
    handles it well in practice (regex.md rule 6).
    """
    leading = 0
    for ch in line:
        if ch != " ":
            break
        leading += 1
    if leading > 3:
        return None  # 4+ spaces: indented code block, not a fence
    if leading >= len(line):
        return None
    fence_char = line[leading]
    if fence_char not in ("`", "~"):
        return None
    run_end = leading
    while run_end < len(line) and line[run_end] == fence_char:
        run_end += 1
    run_len = run_end - leading
    if run_len < 3:
        return None
    rest = line[run_end:]
    # CommonMark §4.5: backtick info strings cannot contain backticks
    # (otherwise the run would be ambiguous with the closing fence).
    # Tilde fences have no such restriction.
    if fence_char == "`" and "`" in rest:
        return None
    return fence_char, run_len, rest


def _push_heading_onto_path(current_path: List[str], heading_line: str, level: int) -> None:
    """Pop deeper-or-equal levels off ``current_path`` then push ``heading_line``."""
    while current_path:
        last_level = len(current_path[-1].split(maxsplit=1)[0])
        if last_level >= level:
            current_path.pop()
        else:
            break
    current_path.append(heading_line)


def _trim_blank_edges(lines: List[str]) -> List[str]:
    """Strip leading/trailing whitespace-only lines; preserve internal blanks."""
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _advance_fence_state(
    open_fence: Optional[Tuple[str, int]],
    fence_info: Tuple[str, int, str],
) -> Optional[Tuple[str, int]]:
    """Return the new ``open_fence`` after seeing a fence-shaped line.

    Per CommonMark §4.5 the closing fence must (1) use the same character
    as the opener, (2) have at least as many fence characters as the
    opener, and (3) carry no info string. A line that fails any of those
    while inside a block is treated as content — caller still appends it
    but the state does not toggle.
    """
    fence_char, fence_len, rest = fence_info
    if open_fence is None:
        # Opener — info string is allowed; record (char, length).
        return (fence_char, fence_len)
    if open_fence[0] == fence_char and fence_len >= open_fence[1] and not rest.strip():
        # Valid close: same char, ≥ as many chars, no info string.
        return None
    # Mismatched char, too short, or has an info string — block stays open.
    return open_fence


def _flush_section(
    sections: List[Tuple[List[str], List[str]]],
    current_path: List[str],
    current_lines: List[str],
    seen_heading: bool,
) -> None:
    """Append the current section to ``sections`` if it carries any content."""
    if seen_heading or current_lines:
        sections.append((list(current_path), current_lines))


def _render_sections(
    sections: List[Tuple[List[str], List[str]]],
) -> List[Tuple[List[str], str]]:
    """Trim blank edges and join each section's body lines into a single string."""
    rendered: List[Tuple[List[str], str]] = []
    for path, body_lines in sections:
        trimmed = _trim_blank_edges(body_lines)
        if trimmed:
            rendered.append((path, "\n".join(trimmed)))
    return rendered


def _markdown_sections(text: str) -> List[Tuple[List[str], str]]:
    """Split a markdown document into ``(heading_path, body)`` sections.

    A section starts at any heading line (``# H1`` … ``###### H6``) and
    runs until the next same-or-higher-level heading. Each section's
    ``heading_path`` is the chain of enclosing headings — e.g. a
    section opened by ``## Background`` inside a document whose first
    heading was ``# Project Notes`` carries
    ``["# Project Notes", "## Background"]``.

    Code fences (``` ``` ``` or ``~~~``) are detected *line-wise* so a
    heading-shaped line **inside** a code block is not interpreted as a
    heading. The CommonMark §4.5 closing-fence rules live in
    :func:`_advance_fence_state`.
    """
    sections: List[Tuple[List[str], List[str]]] = []
    current_path: List[str] = []
    current_lines: List[str] = []
    open_fence: Optional[Tuple[str, int]] = None
    seen_heading = False

    for line in text.splitlines():
        fence_info = _parse_md_fence(line)
        if fence_info is not None:
            open_fence = _advance_fence_state(open_fence, fence_info)
            current_lines.append(line)
            continue
        if open_fence is not None:
            current_lines.append(line)
            continue
        match = _MARKDOWN_HEADING_PATTERN.match(line)
        if not match:
            current_lines.append(line)
            continue
        _flush_section(sections, current_path, current_lines, seen_heading)
        heading_text = match.group(0).strip()
        _push_heading_onto_path(current_path, heading_text, level=len(match.group(1)))
        current_lines = [heading_text]
        seen_heading = True

    _flush_section(sections, current_path, current_lines, seen_heading)
    return _render_sections(sections)


def _heading_breadcrumb(path: List[str], current_heading: str) -> str:
    """Render the breadcrumb that precedes a section's body in a chunk.

    When a section's heading itself is in ``path[-1]``, including the
    full path with the section heading at the end would print the
    heading twice. ``current_heading`` is the section's own heading
    line; we prepend the *parents* (``path[:-1]``) and let the section
    body lead with its own heading.
    """
    parents = [p for p in path if p != current_heading]
    if not parents:
        return ""
    return "\n".join(parents) + "\n\n"


def _chunk_markdown(text: str, max_chunk_size: int) -> Iterable[str]:
    """Heading-aware chunker — Phase 12 third strategy.

    Each chunk is one or more contiguous markdown sections (heading +
    body) packed greedily up to ``max_chunk_size`` characters. Section
    boundaries are heading lines; chunks never split mid-section unless
    a single section already exceeds the cap, in which case it is
    emitted whole (mirrors :func:`_chunk_paragraph`'s "long-paragraph
    on its own" rule).

    Each chunk inlines its enclosing-heading **breadcrumb** at the top
    so the SFT loss sees the document context. For example:

        # Project Notes / ## Background

        ## Background

        body of the section here…

    Code-fenced blocks (``` ``` ```) are kept atomic — never split
    mid-block — because slicing through a code fence produces invalid
    markdown that confuses downstream tokenisers and the model.

    Markdown mode emits **non-overlapping** chunks by design — sections
    are the unit of indivisibility, and an overlap would slice
    mid-section and break the heading-breadcrumb invariant. Use
    ``--strategy sliding`` if you need overlapping windows.
    """
    if max_chunk_size <= 0:
        raise ValueError(_CHUNK_SIZE_POSITIVE_MSG)
    sections = _markdown_sections(text)
    if not sections:
        return

    current: List[str] = []
    current_len = 0
    for path, body in sections:
        breadcrumb = _heading_breadcrumb(path, path[-1] if path else "")
        section_block = breadcrumb + body if breadcrumb else body
        section_len = len(section_block)
        if current_len + section_len + 2 <= max_chunk_size or not current:
            current.append(section_block)
            current_len += section_len + 2
        else:
            yield "\n\n".join(current)
            current = [section_block]
            current_len = section_len
    if current:
        yield "\n\n".join(current)


def _build_markdown_section_blocks(sections: List[Tuple[List[str], str]]) -> List[str]:
    """Render each (heading-path, body) pair into a single text block.

    Pre-built so the tokenizer can be called once on the full list rather
    than per-section. Empty heading paths fall back to the body alone.
    """
    blocks: List[str] = []
    for path, body in sections:
        breadcrumb = _heading_breadcrumb(path, path[-1] if path else "")
        blocks.append(breadcrumb + body if breadcrumb else body)
    return blocks


def _count_section_tokens(section_blocks: List[str], tokenizer: Any) -> List[int]:
    """Return per-block token counts using a single batch call when possible.

    Per-section ``tokenizer.encode()`` called N times used to dominate the
    cost for long documents; the batch call amortises Python-level overhead
    and lets the tokenizer process the list in one pass.  Falls back to
    per-block encoding *only* for tokenizers that genuinely don't accept
    list input — any other failure is re-raised so real bugs (corrupted
    input, OOM, etc.) aren't silently masked behind the slow path.
    """
    try:
        batch_enc = tokenizer(section_blocks, add_special_tokens=False)
    except (TypeError, ValueError):
        # Old / minimal tokenizer APIs raise TypeError ("expected str, got
        # list") or ValueError ("text input must be of type str") when
        # handed a list. That's the documented batching-unsupported
        # signal — fall back to per-block encoding.
        return [len(tokenizer.encode(block, add_special_tokens=False)) for block in section_blocks]
    # Validate the result shape before trusting it: HuggingFace tokenizers
    # return a BatchEncoding (mapping) whose ``input_ids`` is a list of
    # token-id sequences with length == len(section_blocks).  A different
    # shape means the tokenizer accepted the call but produced something
    # we cannot interpret — re-raise instead of silently miscounting.
    try:
        ids_list = batch_enc["input_ids"]
    except (KeyError, TypeError) as e:
        raise TypeError(
            f"Batch tokenizer returned an unexpected shape (no 'input_ids' key): {type(batch_enc).__name__}"
        ) from e
    if not hasattr(ids_list, "__len__") or len(ids_list) != len(section_blocks):
        raise ValueError(
            f"Batch tokenizer returned {len(ids_list) if hasattr(ids_list, '__len__') else '?'} "
            f"sequences for {len(section_blocks)} input blocks; refusing to silently miscount."
        )
    return [len(ids) for ids in ids_list]


def _chunk_markdown_tokens(text: str, max_tokens: int, tokenizer: Any) -> Iterable[str]:
    """Token-aware twin of :func:`_chunk_markdown`. Same semantics, token cap.

    Like the character-mode twin, this strategy is **non-overlapping** —
    section boundaries are atomic. Non-zero ``overlap_tokens`` is
    rejected with a ``ValueError`` by :func:`_strategy_dispatch_tokens`
    *before* this function is called (see ``_MARKDOWN_OVERLAP_UNSUPPORTED_MSG``);
    callers therefore never reach this body with a meaningful overlap.
    """
    if max_tokens <= 0:
        raise ValueError(_CHUNK_TOKENS_POSITIVE_MSG)
    sections = _markdown_sections(text)
    if not sections:
        return

    sep_tokens = len(tokenizer.encode("\n\n", add_special_tokens=False))
    section_blocks = _build_markdown_section_blocks(sections)
    section_token_counts = _count_section_tokens(section_blocks, tokenizer)

    current: List[str] = []
    current_tokens = 0
    for section_block, section_tokens in zip(section_blocks, section_token_counts):
        cost = section_tokens + (sep_tokens if current else 0)
        if current_tokens + cost <= max_tokens or not current:
            current.append(section_block)
            current_tokens += cost
        else:
            yield "\n\n".join(current)
            current = [section_block]
            current_tokens = section_tokens
    if current:
        yield "\n\n".join(current)


def _strategy_dispatch(strategy: str, text: str, chunk_size: int, overlap: int) -> Iterable[str]:
    if strategy == "sliding":
        return _chunk_sliding(text, chunk_size, overlap)
    if strategy == "paragraph":
        return _chunk_paragraph(text, chunk_size)
    if strategy == "markdown":
        if overlap:
            raise ValueError(_MARKDOWN_OVERLAP_UNSUPPORTED_MSG)
        return _chunk_markdown(text, chunk_size)
    if strategy == "semantic":
        return _chunk_semantic(text, chunk_size)
    raise ValueError(f"Unknown chunking strategy '{strategy}'. Choose from: {', '.join(CHUNK_STRATEGIES)}")


# ---------------------------------------------------------------------------
# Phase 11.5: token-aware chunking (--chunk-tokens)
# ---------------------------------------------------------------------------


def _chunk_sliding_tokens(text: str, n_tokens: int, overlap_tokens: int, tokenizer: Any) -> Iterable[str]:
    """Sliding window measured in tokenizer tokens, not characters.

    Mirrors :func:`_chunk_sliding` (overlap clamped to half of the window
    to prevent quadratic chunk explosion) but tokens are the unit so the
    output sizes line up with ``model.max_length`` exactly.
    """
    if n_tokens <= 0:
        raise ValueError(_CHUNK_TOKENS_POSITIVE_MSG)
    if overlap_tokens < 0 or overlap_tokens >= n_tokens:
        raise ValueError("overlap_tokens must be in [0, chunk_tokens)")
    if overlap_tokens > n_tokens // 2:
        raise ValueError(
            f"overlap_tokens ({overlap_tokens}) must be at most chunk_tokens // 2 ({n_tokens // 2}) "
            "to avoid quadratic chunk count. Reduce --overlap-tokens or increase --chunk-tokens."
        )
    encoded = tokenizer.encode(text, add_special_tokens=False)
    if not encoded:
        return
    step = n_tokens - overlap_tokens
    for start in range(0, len(encoded), step):
        ids = encoded[start : start + n_tokens]
        if not ids:
            return
        decoded = tokenizer.decode(ids, skip_special_tokens=True)
        if decoded.strip():
            yield decoded
        if start + n_tokens >= len(encoded):
            return


def _chunk_paragraph_tokens(text: str, max_tokens: int, tokenizer: Any) -> Iterable[str]:
    """Greedy paragraph packer with a token-count cap.

    Same semantics as :func:`_chunk_paragraph` (paragraphs are the unit
    of indivisibility), but the soft cap is measured in tokens so the
    emitted chunks can't blow past ``model.max_length``. Paragraphs that
    exceed ``max_tokens`` on their own are still emitted whole — better
    than mid-sentence splits.

    Paragraph mode is **non-overlapping by design**. Sliding token-overlap
    would slice mid-paragraph and defeat the boundary-preservation
    invariant. Use ``--strategy sliding`` (with ``--overlap-tokens``) when
    overlap is required; the CLI logs an info note when ``overlap_tokens``
    is set alongside ``--strategy paragraph``.

    The ``"\\n\\n"`` separator is included in the budget — most BPE
    tokenizers map it to 1–2 tokens, so a chunk packed near the cap can
    overflow by ~tens of tokens without this accounting. ``sep_tokens``
    is computed once via the tokenizer and added when joining.
    """
    if max_tokens <= 0:
        raise ValueError(_CHUNK_TOKENS_POSITIVE_MSG)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        return

    sep_tokens = len(tokenizer.encode("\n\n", add_special_tokens=False))
    # Single batch tokenize call (closure F-performance-103) instead of
    # re-encoding each paragraph in the loop. Reuses the markdown-chunker
    # helper which already handles old/minimal tokenizers that don't accept
    # list input.
    paragraph_token_counts = _count_section_tokens(paragraphs, tokenizer)

    current: List[str] = []
    current_tokens = 0
    for para, para_tokens in zip(paragraphs, paragraph_token_counts):
        # When we already have packed paragraphs, joining adds a "\n\n"
        # which costs ``sep_tokens`` on top of the new paragraph itself.
        cost = para_tokens + (sep_tokens if current else 0)
        if current_tokens + cost <= max_tokens or not current:
            current.append(para)
            current_tokens += cost
        else:
            yield "\n\n".join(current)
            current = [para]
            current_tokens = para_tokens
    if current:
        yield "\n\n".join(current)


def _strategy_dispatch_tokens(
    strategy: str,
    text: str,
    *,
    chunk_tokens: int,
    overlap_tokens: int,
    tokenizer: Any,
) -> Iterable[str]:
    if strategy == "sliding":
        return _chunk_sliding_tokens(text, chunk_tokens, overlap_tokens, tokenizer)
    if strategy == "paragraph":
        return _chunk_paragraph_tokens(text, chunk_tokens, tokenizer)
    if strategy == "markdown":
        if overlap_tokens:
            raise ValueError(_MARKDOWN_OVERLAP_UNSUPPORTED_MSG)
        return _chunk_markdown_tokens(text, chunk_tokens, tokenizer)
    if strategy == "semantic":
        return _chunk_semantic(text, chunk_tokens)
    raise ValueError(f"Unknown chunking strategy '{strategy}'. Choose from: {', '.join(CHUNK_STRATEGIES)}")


def _load_tokenizer(model_name: str) -> Any:
    """Resolve ``model_name`` to an HF :class:`AutoTokenizer`.

    Lazy import keeps the ingestion module import-time small for users
    who never touch token-aware chunking. ``trust_remote_code`` is
    intentionally OFF — token-aware chunking should never need a custom
    tokenizer class, and turning it on here would let arbitrary HF Hub
    repos run code at ingestion time without a config audit.
    """
    from transformers import AutoTokenizer  # type: ignore

    return AutoTokenizer.from_pretrained(model_name, use_fast=True, trust_remote_code=False)


# ---------------------------------------------------------------------------
# File discovery + ingestion entry point
# ---------------------------------------------------------------------------


def _iter_input_files(input_path: Path, recursive: bool) -> Iterable[Path]:
    if input_path.is_file():
        yield input_path
        return
    if not input_path.is_dir():
        raise FileNotFoundError(f"Input path does not exist: {input_path}")
    pattern = "**/*" if recursive else "*"
    for entry in sorted(input_path.glob(pattern)):
        if entry.is_file() and entry.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield entry


def _select_extractor(path: Path) -> Optional[Callable[[Path], str]]:
    return _EXTRACTORS.get(path.suffix.lower())


def list_supported_formats() -> List[str]:
    """Return the list of file extensions ingest_path can handle."""
    return list(SUPPORTED_EXTENSIONS)


def describe_strategies() -> List[Tuple[str, str]]:
    """Return ``(name, one-line description)`` for every chunking strategy."""
    return [
        ("sliding", "Fixed-size character window with overlap. Predictable, coarse."),
        ("paragraph", "Greedy paragraph packer; never splits a paragraph. Preserves boundaries."),
        (
            "markdown",
            "Heading-aware splitter (Phase 12). Chunks at # / ## / ### boundaries, "
            "keeps code-fenced blocks atomic, inlines a heading breadcrumb so SFT "
            "loss sees document context.",
        ),
        ("semantic", "Embedding-clustered (NotImplementedError today; planned for a follow-up phase)."),
    ]


@dataclass
class _FileOutcome:
    """Per-file aggregate emitted by :func:`_process_one_file`."""

    chunks_written: int = 0
    chars_written: int = 0
    file_processed: bool = False
    file_skipped: bool = False
    extension: Optional[str] = None
    pii_counts: Dict[str, int] = field(default_factory=dict)
    secrets_counts: Dict[str, int] = field(default_factory=dict)


def _extract_text_for_ingest(
    fpath: Path,
    extractor: Callable[..., str],
    ctx: _ExtractContext,
) -> Optional[str]:
    """Run ``extractor`` against ``fpath``; return ``None`` to mean "skip this file".

    Phase 15: threads the shared :class:`_ExtractContext` into every
    extractor (uniform signature ``(path, *, ctx)``) and applies the
    post-extract normalisation pipeline (glyph normalise → script sanity
    → strip-patterns → strip-URLs) before returning.

    ``ImportError`` re-propagates (missing optional extra is a *runtime*
    failure of the dispatched feature, not a per-file skip). All other
    exceptions log a warning and signal a skip — consistent with the
    silently-tolerant per-file model the CLI relies on.
    """
    try:
        text = extractor(fpath, ctx=ctx)
    except (ImportError, IngestParameterError):
        # ImportError: missing optional extra is a *runtime* failure of
        # the dispatched feature — propagate so the CLI dispatcher can
        # translate to EXIT_TRAINING_ERROR with a clear install-hint.
        # IngestParameterError (Phase 15): operator-supplied parameter
        # mistake — propagate so the run aborts with EXIT_CONFIG_ERROR
        # instead of silently soft-failing per-file.
        raise
    except Exception as exc:  # noqa: BLE001 — best-effort: per-file dispatcher catch — each format-specific extractor (PDF/DOCX/EPUB/TXT/MD) raises its own typed ValueError above + a wide tail of corpus-data-driven failures; per-file soft-fail with skip-and-continue keeps a multi-thousand-file corpus ingest running.  ImportError stays narrow above so missing-extra failures still propagate.  # NOSONAR
        logger.warning("Skipping '%s' (extraction failed): %s", fpath, exc)
        return None
    if not text:
        return text
    return _post_extract_normalise(text, ctx, fpath)


def _select_chunks_iter(
    text: str,
    *,
    strategy: str,
    chunk_size: int,
    overlap: int,
    chunk_tokens: Optional[int],
    overlap_tokens: int,
    tokenizer: Any,
) -> Iterable[str]:
    """Pick the token-aware or character-based chunker based on flags."""
    if chunk_tokens and tokenizer is not None:
        return _strategy_dispatch_tokens(
            strategy,
            text,
            chunk_tokens=chunk_tokens,
            overlap_tokens=overlap_tokens,
            tokenizer=tokenizer,
        )
    return _strategy_dispatch(strategy, text, chunk_size, overlap)


def _emit_chunk(
    payload: str,
    out_fh: Any,
    outcome: _FileOutcome,
    mask_pii: Optional[Callable[..., Any]],
    mask_secrets: Optional[Callable[..., Any]] = None,
    sampler: Optional[Callable[[str], None]] = None,
) -> None:
    """Mask (optional), serialise, and write one chunk; update outcome counters.

    Mask order is **secrets first, PII second** when both are enabled.
    This is a **defensive ordering**: today's regex sets do not produce
    measurable cross-detector overlap on the test fixtures we ship, but
    secrets are higher-severity than PII (a leaked AWS key in training
    data is unrecoverable; a phone number is recoverable via opt-out
    flows), so when ordering matters at all it should favour the secrets
    pass. Future PII / secret regex additions could legitimately overlap
    (e.g. an Azure connection-string substring resembling an IBAN); this
    ordering future-proofs against that case without operator action.

    Round-2 review (S-3): the quality pre-signal previously read each
    written line back via ``json.loads`` to recover the chunk text, an
    avoidable per-chunk round-trip on large corpora. The ``sampler``
    kwarg now hands the post-mask chunk string straight to the
    pre-signal collector, bypassing the JSON parse entirely.
    """
    if mask_secrets is not None:
        payload, secret_counts = mask_secrets(payload, return_counts=True)
        for kind, count in secret_counts.items():
            outcome.secrets_counts[kind] = outcome.secrets_counts.get(kind, 0) + count
    if mask_pii is not None:
        # Get the masked text + per-type counts in a single pass. Counting
        # via detect_pii beforehand would double-count spans matched by
        # multiple patterns; mask_pii's own first-match-wins precedence
        # gives the truthful count.
        payload, redaction_counts = mask_pii(payload, return_counts=True)
        for kind, count in redaction_counts.items():
            outcome.pii_counts[kind] = outcome.pii_counts.get(kind, 0) + count
    if sampler is not None:
        sampler(payload)
    out_fh.write(json.dumps({"text": payload}, ensure_ascii=False) + "\n")
    outcome.chunks_written += 1
    outcome.chars_written += len(payload)


def _process_one_file(
    fpath: Path,
    out_fh: Any,
    *,
    strategy: str,
    chunk_size: int,
    overlap: int,
    mask_pii: Optional[Callable[..., Any]],
    mask_secrets: Optional[Callable[..., Any]] = None,
    chunk_tokens: Optional[int] = None,
    overlap_tokens: int = 0,
    tokenizer: Any = None,
    extract_ctx: Optional[_ExtractContext] = None,
    sampler: Optional[Callable[[str], None]] = None,
) -> _FileOutcome:
    """Extract → chunk → optionally mask → emit JSONL for a single file.

    ``chunk_tokens`` is honoured when set: token-aware chunking takes over
    via :func:`_strategy_dispatch_tokens`, and ``chunk_size`` becomes a
    fallback only when the tokenizer is missing. ``extract_ctx`` carries
    the Phase 15 extraction options + side-output state (header/footer
    dedup counters, script-sanity reports, front-matter dropped pages,
    etc.) so they roll up into the run-level structured notes.

    ImportError propagates (missing optional extra is not a per-file skip).
    Any other extraction failure is logged + counted as a skip.
    """
    if extract_ctx is None:
        extract_ctx = _ExtractContext()
    extractor = _select_extractor(fpath)
    if extractor is None:
        return _FileOutcome(file_skipped=True)

    text = _extract_text_for_ingest(fpath, extractor, extract_ctx)
    if text is None:
        return _FileOutcome(file_skipped=True)
    if not text or not text.strip():
        logger.warning("Skipping '%s' (no extractable text).", fpath)
        return _FileOutcome(file_skipped=True)

    outcome = _FileOutcome(file_processed=True, extension=fpath.suffix.lower())
    chunks_iter = _select_chunks_iter(
        text,
        strategy=strategy,
        chunk_size=chunk_size,
        overlap=overlap,
        chunk_tokens=chunk_tokens,
        overlap_tokens=overlap_tokens,
        tokenizer=tokenizer,
    )
    for chunk in chunks_iter:
        payload = chunk.strip()
        if not payload:
            continue
        _emit_chunk(payload, out_fh, outcome, mask_pii, mask_secrets=mask_secrets, sampler=sampler)
    return outcome


DEFAULT_CHUNK_SIZE: int = 2048
"""Public default for the character-based chunk-size cap.

Exposed so the CLI default and the library default share a single source
of truth, and so the "did the operator pass --chunk-size explicitly?"
detection in :func:`ingest_path` is not a magic-number compare.
"""


DEFAULT_SLIDING_OVERLAP: int = 200


def ingest_path(  # NOSONAR python:S107 — every kwarg is a documented operator-facing knob; collapsing them into a config dataclass would be a breaking API change for v0.5 callers that already pass these by name (per `forgelm.__all__` / `__api_version__` contract).
    input_path: str,
    *,
    output_path: str,
    chunk_size: Optional[int] = None,
    overlap: Optional[int] = None,
    strategy: str = "paragraph",
    recursive: bool = False,
    pii_mask: bool = False,
    secrets_mask: bool = False,
    encoding: str = "utf-8",
    chunk_tokens: Optional[int] = None,
    overlap_tokens: int = 0,
    tokenizer: Optional[str] = None,
    # ---- Phase 15 knobs -------------------------------------------------
    language_hint: Optional[str] = None,
    script_sanity_threshold: float = _DEFAULT_SCRIPT_SANITY_THRESHOLD,
    # ``None`` triggers the C-2 auto-derive: when ``language_hint=="tr"``
    # the profile is "turkish"; otherwise "none". An explicit value wins.
    normalise_profile: Optional[str] = None,
    keep_md_frontmatter: bool = False,
    epub_skip_frontmatter: bool = True,
    keep_frontmatter: bool = False,
    page_range: Optional[Tuple[int, int]] = None,
    strip_patterns: Optional[List[str]] = None,
    strip_pattern_timeout: Optional[int] = 5,
    strip_urls: str = "keep",
    quality_presignal: bool = True,
) -> IngestionResult:
    """Ingest a single file or a directory tree into a SFT-compatible JSONL.

    Args:
        input_path: File or directory to ingest.
        output_path: Where to write the ``.jsonl`` output. Parents are created.
        chunk_size: Soft size cap per chunk (characters). ``None`` means
            "use the library default" (:data:`DEFAULT_CHUNK_SIZE`). When
            ``chunk_tokens`` is set the value is ignored, and a warning is
            logged only when the operator passed an *explicit* ``chunk_size``
            so stale CLI invocations are visible without spamming the
            common case.
        overlap: Overlap window for the sliding strategy in characters.
            Must satisfy both ``overlap < chunk_size`` AND
            ``overlap <= chunk_size // 2``. The half-chunk cap prevents
            quadratic chunk explosion. ``_chunk_sliding`` raises
            ``ValueError`` when either bound is violated.
        strategy: One of ``sliding`` / ``paragraph`` / ``markdown`` / ``semantic``.
            Phase 12 added ``markdown`` — heading-aware splitter that
            preserves heading hierarchy and code-block boundaries.
        recursive: When ``input_path`` is a directory, walk subdirectories too.
        pii_mask: Replace detected PII spans with ``[REDACTED]`` before writing.
        secrets_mask: Phase 12 — replace detected credential/secret spans
            (AWS / GitHub / Slack / OpenAI / Google / JWT / OpenSSH / PGP /
            Azure storage) with ``[REDACTED-SECRET]`` before chunks land
            in the JSONL. Runs *before* PII masking when both are enabled
            so secrets matched by both detectors are scrubbed under the
            stronger label first.
        encoding: Output encoding (default UTF-8).
        chunk_tokens: Phase 11.5 token-aware mode. When set, chunks are
            sized against the supplied ``tokenizer`` (in tokens), not
            characters. Use this when the operator's downstream model has
            a hard ``max_length`` budget and char-based sizing keeps
            tripping it.
        overlap_tokens: Sliding-window overlap measured in tokens. Same
            half-window cap as the character-based ``overlap``. Ignored
            when ``strategy="paragraph"`` (paragraph chunks are
            non-overlapping by design); a warning is logged in that case.
        tokenizer: HuggingFace model name resolved via :class:`AutoTokenizer`.
            Required when ``chunk_tokens`` is set; ignored otherwise.

    Returns:
        :class:`IngestionResult` summarizing the run.

    Raises:
        FileNotFoundError: ``input_path`` does not exist.
        ValueError: invalid chunking parameters or ``chunk_tokens`` set
            without a ``tokenizer``.
        ImportError: optional extras not installed for the format being ingested.
    """
    src = Path(input_path).expanduser().resolve()
    dst = Path(output_path).expanduser().resolve()
    dst.parent.mkdir(parents=True, exist_ok=True)

    if strip_urls not in ("keep", "mask", "strip"):
        raise ValueError(f"strip_urls must be one of 'keep' / 'mask' / 'strip', got {strip_urls!r}.")
    if page_range is not None and (not isinstance(page_range, tuple) or len(page_range) != 2):
        raise ValueError("page_range must be a (start, end) tuple of 1-indexed inclusive page numbers.")

    # Phase 15 Wave 2 Task 11: compile + validate operator-supplied
    # --strip-pattern values up-front so a malformed regex aborts the
    # run BEFORE any I/O.  The CLI dispatcher catches StripPatternError
    # (subclass of ValueError) and exits with EXIT_CONFIG_ERROR.
    compiled_strip_patterns: List[Tuple[str, "re.Pattern[str]"]] = []
    if strip_patterns:
        from ._strip_pattern import compile_strip_patterns

        compiled_strip_patterns = compile_strip_patterns(strip_patterns)

    # Phase 15 round-2 (C-2) library-level profile derivation.  When the
    # caller passes ``normalise_profile=None`` we derive "turkish" from
    # ``language_hint == "tr"`` and "none" otherwise.  Explicit kwarg
    # values (including the literal ``"none"``) bypass derivation so
    # operator intent is honoured.  The CLI dispatcher mirrors this
    # logic in ``_resolve_normalise_profile`` for parity.
    if normalise_profile is None:
        normalise_profile = "turkish" if (language_hint or "").lower() == "tr" else "none"

    chunk_size_explicit = chunk_size is not None
    effective_chunk_size = chunk_size if chunk_size_explicit else DEFAULT_CHUNK_SIZE

    # Resolve ``overlap``: only the sliding strategy needs a non-zero default.
    # Paragraph and markdown are non-overlapping by design — we keep
    # ``overlap=None`` flowing into the dispatcher so the default doesn't
    # spuriously trigger the markdown overlap-rejected validator. For
    # sliding, clamp the implicit default to ``effective_chunk_size // 2``
    # so a small ``--chunk-size`` (e.g. 300) doesn't trip ``_chunk_sliding``'s
    # "overlap > chunk_size // 2" guard with the default 200, which would
    # surface as a confusing error for a knob the user didn't even set.
    if overlap is None:
        if strategy == "sliding":
            overlap = min(DEFAULT_SLIDING_OVERLAP, max(0, effective_chunk_size // 2))
        else:
            overlap = 0

    if pii_mask:
        # Lazy import: PII helpers live in data_audit.py; we don't want to
        # pay the audit module's import cost when masking is off.
        from .data_audit import mask_pii as _mask_pii

        mask_pii_fn: Optional[Callable[..., Any]] = _mask_pii
    else:
        mask_pii_fn = None

    if secrets_mask:
        # Phase 12: same lazy-import pattern as ``mask_pii``. The secrets
        # detector lives in ``data_audit`` because the same regex set is
        # used by ``forgelm audit --secrets-mask`` (audit-side reporting)
        # and by ingest-side scrubbing here.
        from .data_audit import mask_secrets as _mask_secrets

        mask_secrets_fn: Optional[Callable[..., Any]] = _mask_secrets
    else:
        mask_secrets_fn = None

    tokenizer_obj: Any = None
    if chunk_tokens is not None:
        if not tokenizer:
            raise ValueError(
                "--chunk-tokens requires --tokenizer MODEL_NAME so the audit can size chunks against the right vocab."
            )
        if chunk_size_explicit:
            logger.warning(
                "Token-aware mode active (--chunk-tokens=%d): --chunk-size=%d is ignored.",
                chunk_tokens,
                effective_chunk_size,
            )
        tokenizer_obj = _load_tokenizer(tokenizer)
    if strategy == "paragraph" and overlap_tokens > 0:
        # Paragraph mode is intentionally non-overlapping (paragraphs are
        # the unit of indivisibility). Surfacing this as a one-line note
        # keeps the operator from silently losing their --overlap-tokens.
        logger.info(
            "--overlap-tokens=%d is ignored for strategy='paragraph' "
            "(paragraph chunks don't overlap by design). Use --strategy sliding for token-overlap.",
            overlap_tokens,
        )

    files = list(_iter_input_files(src, recursive))
    if not files:
        raise FileNotFoundError(f"No supported files found at '{src}' (extensions: {', '.join(SUPPORTED_EXTENSIONS)}).")

    chunk_count = 0
    files_processed = 0
    files_skipped = 0
    total_chars = 0
    format_counts: Dict[str, int] = {}
    pii_redaction_counts: Dict[str, int] = {}
    secrets_redaction_counts: Dict[str, int] = {}
    notes: List[str] = []
    pdf_dedup_state: Dict[str, int] = {}

    extract_ctx = _ExtractContext(
        dedup_state=pdf_dedup_state,
        page_range=page_range,
        normalise_profile=normalise_profile,
        language_hint=language_hint,
        script_sanity_threshold=script_sanity_threshold,
        keep_md_frontmatter=keep_md_frontmatter,
        epub_skip_frontmatter=epub_skip_frontmatter,
        keep_frontmatter=keep_frontmatter,
        strip_patterns=compiled_strip_patterns,
        strip_pattern_timeout=strip_pattern_timeout,
        strip_urls_mode=strip_urls,
    )

    # Phase 15 Task 4: collect each chunk's text on the fly so the
    # end-of-run quality pre-signal can compute its three cheap checks
    # without re-reading the JSONL. Round-2 review (S-3): we now pass
    # the sampler straight into ``_emit_chunk`` instead of wrapping the
    # file handle and re-parsing the JSON payload back out of every
    # written line.
    sampled_chunks: List[str] = []
    quality_sample_limit = 5000  # keep memory bounded on huge corpora

    sampler: Optional[Callable[[str], None]]
    if quality_presignal:

        def _maybe_sample(payload: str) -> None:
            if len(sampled_chunks) < quality_sample_limit:
                sampled_chunks.append(payload)

        sampler = _maybe_sample
    else:
        sampler = None

    # newline="\n" pins LF on Windows. JSONL Files spec requires LF, and
    # piping through tooling (jq -c, wc -l, downstream HF dataset loaders)
    # avoids CRLF surprises. Linux/macOS default is already LF.
    with open(dst, "w", encoding=encoding, newline="\n") as out_fh:
        for fpath in files:
            outcome = _process_one_file(
                fpath,
                out_fh,
                strategy=strategy,
                chunk_size=effective_chunk_size,
                overlap=overlap,
                mask_pii=mask_pii_fn,
                mask_secrets=mask_secrets_fn,
                chunk_tokens=chunk_tokens,
                overlap_tokens=overlap_tokens,
                tokenizer=tokenizer_obj,
                extract_ctx=extract_ctx,
                sampler=sampler,
            )
            chunk_count += outcome.chunks_written
            total_chars += outcome.chars_written
            if outcome.file_processed:
                files_processed += 1
                if outcome.extension:
                    format_counts[outcome.extension] = format_counts.get(outcome.extension, 0) + 1
            if outcome.file_skipped:
                files_skipped += 1
            for kind, count in outcome.pii_counts.items():
                pii_redaction_counts[kind] = pii_redaction_counts.get(kind, 0) + count
            for kind, count in outcome.secrets_counts.items():
                secrets_redaction_counts[kind] = secrets_redaction_counts.get(kind, 0) + count

    if files_skipped:
        notes.append(f"skipped {files_skipped} file(s) — see warnings above")
    if pii_mask:
        if pii_redaction_counts:
            redacted_total = sum(pii_redaction_counts.values())
            breakdown = ", ".join(f"{k}={v}" for k, v in sorted(pii_redaction_counts.items()))
            notes.append(f"PII masking redacted {redacted_total} span(s): {breakdown}")
        else:
            notes.append("PII masking enabled — no PII detected in this corpus")
    if secrets_mask:
        if secrets_redaction_counts:
            secret_total = sum(secrets_redaction_counts.values())
            breakdown = ", ".join(f"{k}={v}" for k, v in sorted(secrets_redaction_counts.items()))
            notes.append(f"Secrets masking redacted {secret_total} credential span(s): {breakdown}")
        else:
            notes.append("Secrets masking enabled — no credentials detected in this corpus")
    pdf_lines_stripped = pdf_dedup_state.get("lines_stripped", 0)
    pdf_paragraph_stripped = pdf_dedup_state.get("paragraph_packed_stripped", 0)
    if pdf_lines_stripped:
        notes.append(
            f"PDF header/footer dedup stripped {pdf_lines_stripped} repeated line(s) "
            "(reduces audit near-duplicate noise)."
        )
    if pdf_paragraph_stripped:
        notes.append(f"PDF post-pack dedup stripped {pdf_paragraph_stripped} survivor header line(s).")

    # Phase 15 Task 4: quality pre-signal at end of run.
    quality_presignal_payload: Optional[Dict[str, Any]] = None
    if quality_presignal and sampled_chunks:
        quality_presignal_payload = _compute_quality_presignal(sampled_chunks)
        flagged = quality_presignal_payload["samples_flagged"]
        evaluated = quality_presignal_payload["samples_evaluated"]
        if flagged:
            notes.append(
                f"{flagged}/{evaluated} chunks below ingestion quality threshold. "
                "Run `forgelm audit <output>` for detail."
            )
            # The audit guide names the same call shape, so the operator
            # can grep this string to find the doc page. Emitting via
            # logger.warning makes it visible in CI logs regardless of
            # --output-format json/text.
            logger.warning(
                "[WARN] %d/%d chunks below ingestion quality threshold. Run `forgelm audit %s` for detail.",
                flagged,
                evaluated,
                dst,
            )

    # Phase 15 Task 2: aggregate script-sanity reports.
    from ._script_sanity import report_to_structured as _script_sanity_to_structured

    script_sanity_summary: Optional[Dict[str, Any]] = None
    if extract_ctx.script_sanity_reports:
        script_sanity_summary = _script_sanity_to_structured(extract_ctx.script_sanity_reports)
        triggered_count = script_sanity_summary.get("files_triggered", 0)
        if triggered_count:
            notes.append(f"Script-sanity check fired on {triggered_count} file(s) — see warnings.")

    structured_notes: Dict[str, Any] = {
        "files_processed": files_processed,
        "files_skipped": files_skipped,
        "chunk_count": chunk_count,
        "total_chars": total_chars,
        "strategy": strategy,
        "format_counts": dict(format_counts),
        "pii_redaction_counts": dict(pii_redaction_counts),
        "secrets_redaction_counts": dict(secrets_redaction_counts),
    }
    if pdf_lines_stripped:
        structured_notes["pdf_header_footer_lines_stripped"] = pdf_lines_stripped
    if pdf_paragraph_stripped:
        structured_notes["pdf_paragraph_packed_lines_stripped"] = pdf_paragraph_stripped
    if chunk_tokens is not None:
        structured_notes["chunk_tokens"] = chunk_tokens
        structured_notes["tokenizer"] = tokenizer
    if script_sanity_summary is not None:
        structured_notes["script_sanity_summary"] = script_sanity_summary
    if extract_ctx.frontmatter_dropped_pages:
        structured_notes["frontmatter_pages_dropped"] = sorted(set(extract_ctx.frontmatter_dropped_pages))
    if extract_ctx.strip_pattern_substitutions_total:
        structured_notes["strip_pattern_substitutions"] = extract_ctx.strip_pattern_substitutions_total
        notes.append(f"--strip-pattern removed {extract_ctx.strip_pattern_substitutions_total} span(s).")
    if extract_ctx.urls_handled_total:
        structured_notes["urls_handled"] = extract_ctx.urls_handled_total
        action = {"mask": "masked", "strip": "stripped"}.get(strip_urls, strip_urls)
        notes.append(f"--strip-urls {action} {extract_ctx.urls_handled_total} URL(s).")
    if quality_presignal_payload is not None:
        structured_notes["quality_presignal"] = quality_presignal_payload

    logger.info(
        "ingest: source=%s output=%s files=%d chunks=%d chars=%d strategy=%s",
        src,
        dst,
        files_processed,
        chunk_count,
        total_chars,
        strategy,
    )

    return IngestionResult(
        output_path=dst,
        chunk_count=chunk_count,
        files_processed=files_processed,
        files_skipped=files_skipped,
        total_chars=total_chars,
        format_counts=format_counts,
        pii_redaction_counts=pii_redaction_counts,
        secrets_redaction_counts=secrets_redaction_counts,
        extra_notes=notes,
        notes_structured=structured_notes,
        pdf_header_footer_lines_stripped=pdf_lines_stripped,
        pdf_paragraph_packed_lines_stripped=pdf_paragraph_stripped,
        script_sanity_triggered=(script_sanity_summary or {}).get("files_triggered", 0) if script_sanity_summary else 0,
        strip_pattern_substitutions=extract_ctx.strip_pattern_substitutions_total,
        urls_handled=extract_ctx.urls_handled_total,
        frontmatter_pages_dropped=len(set(extract_ctx.frontmatter_dropped_pages)),
    )


# ---------------------------------------------------------------------------
# Phase 15 Task 4 — end-of-run quality pre-signal helpers
# ---------------------------------------------------------------------------


def _check_alpha_ratio(text: str) -> bool:
    """Phase 15 Task 4 — alpha-ratio cheap-check.

    Returns ``True`` when the chunk fails the check (alpha ratio < 0.70).
    Same threshold as the audit's full quality filter so the pre-signal
    aligns with the deeper diagnostic.
    """
    non_ws = [c for c in text if not c.isspace()]
    if not non_ws:
        return True
    alpha = sum(1 for c in non_ws if c.isalpha()) / len(non_ws)
    return alpha < 0.70


def _check_weird_char_ratio(text: str) -> bool:
    """Phase 15 Task 4 — weird-character cheap-check.

    Returns ``True`` when the chunk has > 5 % "weird" characters (control
    chars + U+FFFD replacement char + isolated PUA glyphs). Catches the
    audit's font-corruption mode without the full Unicode-block sanity
    pass.
    """
    if not text:
        return False
    weird = sum(
        1
        for c in text
        if c == "�"
        or (ord(c) < 0x20 and c not in ("\n", "\r", "\t"))
        or (0xE000 <= ord(c) <= 0xF8FF)
        or (0x0800 <= ord(c) <= 0x097F and not c.isalpha())
    )
    return weird / len(text) > 0.05


def _check_repeated_line_ratio(text: str) -> bool:
    """Phase 15 Task 4 — repeated-line cheap-check.

    Returns ``True`` when ≥ 30 % of a chunk's lines are duplicates.
    Same threshold as the audit's full ``repeated_lines`` check, but
    cheap because we work on already-emitted chunk text instead of
    re-reading the JSONL.
    """
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if len(lines) < 3:
        return False
    counts = Counter(lines)
    repeated = sum(n for ln, n in counts.items() if n >= 2)
    return repeated / len(lines) > 0.30


def _compute_quality_presignal(chunks: List[str]) -> Dict[str, Any]:
    """Apply the three Phase 15 Task 4 cheap checks to a sample of chunks.

    Returns a structured payload that always carries the schema
    callers depend on, even when nothing was flagged. The threshold,
    sample size, and per-check counts are surfaced so downstream
    tooling (audit, CI gates) can decide what to do — we do not block
    the run.
    """
    flagged = 0
    by_check = {"alpha_ratio": 0, "weird_chars": 0, "repeated_lines": 0}
    for chunk in chunks:
        chunk_flagged = False
        if _check_alpha_ratio(chunk):
            by_check["alpha_ratio"] += 1
            chunk_flagged = True
        if _check_weird_char_ratio(chunk):
            by_check["weird_chars"] += 1
            chunk_flagged = True
        if _check_repeated_line_ratio(chunk):
            by_check["repeated_lines"] += 1
            chunk_flagged = True
        if chunk_flagged:
            flagged += 1
    return {
        "samples_evaluated": len(chunks),
        "samples_flagged": flagged,
        "by_check": by_check,
    }


def summarize_result(result: IngestionResult) -> str:
    """Render an :class:`IngestionResult` as a multi-line operator-friendly report."""
    lines = [
        "Ingestion summary",
        f"  Output JSONL  : {result.output_path}",
        f"  Files (in/out): processed={result.files_processed}  skipped={result.files_skipped}",
        f"  Chunks        : {result.chunk_count}",
        f"  Total chars   : {result.total_chars}",
    ]
    if result.format_counts:
        per_fmt = ", ".join(f"{ext.lstrip('.')}={n}" for ext, n in sorted(result.format_counts.items()))
        lines.append(f"  By format     : {per_fmt}")
    for note in result.extra_notes:
        lines.append(f"  Note          : {note}")
    lines.append("")
    lines.append(f"Next: forgelm audit {result.output_path} --output ./audit/")
    lines.append(
        f"Or train: forgelm --config <your.yaml>  (set data.dataset_name_or_path: {os.fspath(result.output_path)})"
    )
    return "\n".join(lines)
