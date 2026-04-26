"""Tests for forgelm.ingestion (Phase 11).

The TXT path and chunking strategies are stdlib-only and are covered without
optional extras. PDF/DOCX/EPUB extractors are skipped when the corresponding
optional dep is missing — this matches how the module behaves in the wild
(``pip install 'forgelm[ingestion]'`` is the install instruction).
"""

from __future__ import annotations

import importlib
import json
from pathlib import Path

import pytest

from forgelm.ingestion import (
    CHUNK_STRATEGIES,
    SUPPORTED_EXTENSIONS,
    IngestionResult,
    _chunk_paragraph,
    _chunk_semantic,
    _chunk_sliding,
    describe_strategies,
    ingest_path,
    list_supported_formats,
    summarize_result,
)

# ---------------------------------------------------------------------------
# Public surface
# ---------------------------------------------------------------------------


class TestPublicSurface:
    def test_supported_formats_listed(self):
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert tuple(list_supported_formats()) == SUPPORTED_EXTENSIONS

    def test_strategies_describe_each_name(self):
        names = {n for n, _desc in describe_strategies()}
        assert names == set(CHUNK_STRATEGIES)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


class TestSlidingChunking:
    def test_no_overlap_partitions_input(self):
        text = "abcdefghij"
        chunks = list(_chunk_sliding(text, chunk_size=4, overlap=0))
        assert chunks == ["abcd", "efgh", "ij"]

    def test_overlap_preserves_context(self):
        text = "abcdefgh"
        chunks = list(_chunk_sliding(text, chunk_size=4, overlap=2))
        # step = 2 → "abcd", "cdef", "efgh", "gh"
        assert chunks[0] == "abcd"
        assert chunks[1] == "cdef"
        assert chunks[2] == "efgh"

    def test_overlap_must_be_less_than_chunk(self):
        with pytest.raises(ValueError, match="overlap"):
            list(_chunk_sliding("abc", chunk_size=2, overlap=2))

    def test_overlap_above_half_chunk_rejected(self):
        # Pathological overlap: chunk_size=200 + overlap=199 would emit
        # ~one chunk per character. Reject up front.
        with pytest.raises(ValueError, match="quadratic"):
            list(_chunk_sliding("x" * 1000, chunk_size=200, overlap=199))

    def test_no_runt_trailing_chunk(self):
        # Bug 5: text=205, chunk=200, overlap=100, step=100 used to emit
        # 3 chunks: [200, 105, 5]. The 5-char tail is a pure prefix of
        # the previous chunk's overlap region — it adds zero new content
        # AND skews near-duplicate / length stats. Early-exit kills it.
        text = "x" * 205
        chunks = list(_chunk_sliding(text, chunk_size=200, overlap=100))
        assert len(chunks) == 2
        assert all(len(c) >= 100 for c in chunks)

    def test_negative_chunk_size_rejected(self):
        with pytest.raises(ValueError, match="chunk_size"):
            list(_chunk_sliding("abc", chunk_size=-1, overlap=0))

    def test_empty_input_is_empty(self):
        assert list(_chunk_sliding("", chunk_size=4, overlap=0)) == []


class TestParagraphChunking:
    def test_packs_short_paragraphs_together(self):
        text = "Para A.\n\nPara B.\n\nPara C."
        chunks = list(_chunk_paragraph(text, max_chunk_size=200))
        # Each para is short; all three should pack into one chunk.
        assert len(chunks) == 1
        assert "Para A." in chunks[0] and "Para C." in chunks[0]

    def test_long_paragraph_emits_alone(self):
        big = "x" * 500
        text = f"short.\n\n{big}\n\nshort again."
        chunks = list(_chunk_paragraph(text, max_chunk_size=100))
        # The 500-char paragraph exceeds the cap; it should land alone.
        assert any(c == big for c in chunks)

    def test_drops_blank_paragraphs(self):
        text = "\n\n\n\nHello.\n\n\n\n"
        chunks = list(_chunk_paragraph(text, max_chunk_size=200))
        assert chunks == ["Hello."]


class TestSemanticChunking:
    def test_raises_not_implemented(self):
        with pytest.raises(NotImplementedError, match="Semantic chunking"):
            list(_chunk_semantic("abc", chunk_size=10))


# ---------------------------------------------------------------------------
# End-to-end ingest_path on TXT (no optional deps required)
# ---------------------------------------------------------------------------


def _write_txt(path: Path, body: str) -> Path:
    path.write_text(body, encoding="utf-8")
    return path


class TestIngestSingleFile:
    def test_writes_jsonl_with_text_column(self, tmp_path):
        src = _write_txt(tmp_path / "doc.txt", "Hello world.\n\nSecond paragraph here.")
        out = tmp_path / "out.jsonl"

        result = ingest_path(str(src), output_path=str(out), strategy="paragraph", chunk_size=200)

        assert isinstance(result, IngestionResult)
        assert out.is_file()
        assert result.files_processed == 1
        assert result.files_skipped == 0
        assert result.chunk_count >= 1
        rows = [json.loads(line) for line in out.read_text(encoding="utf-8").splitlines() if line.strip()]
        assert all(set(r.keys()) == {"text"} for r in rows)
        assert any("Hello world" in r["text"] for r in rows)

    def test_paragraph_strategy_yields_one_chunk_for_small_doc(self, tmp_path):
        src = _write_txt(tmp_path / "small.txt", "A.\n\nB.\n\nC.")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(src), output_path=str(out), strategy="paragraph", chunk_size=200)
        assert result.chunk_count == 1

    def test_sliding_strategy_respects_chunk_size(self, tmp_path):
        src = _write_txt(tmp_path / "big.txt", "x" * 1000)
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(src), output_path=str(out), strategy="sliding", chunk_size=200, overlap=0)
        assert result.chunk_count == 5  # 1000 / 200

    def test_unknown_strategy_raises(self, tmp_path):
        src = _write_txt(tmp_path / "x.txt", "x")
        out = tmp_path / "y.jsonl"
        with pytest.raises(ValueError, match="strategy"):
            ingest_path(str(src), output_path=str(out), strategy="bogus")

    def test_missing_input_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            ingest_path(str(tmp_path / "missing.txt"), output_path=str(tmp_path / "y.jsonl"))

    def test_pii_mask_redacts_emails(self, tmp_path):
        src = _write_txt(tmp_path / "doc.txt", "Contact me at alice@example.com please.")
        out = tmp_path / "redacted.jsonl"
        ingest_path(str(src), output_path=str(out), strategy="paragraph", pii_mask=True)
        body = out.read_text(encoding="utf-8")
        assert "alice@example.com" not in body
        assert "[REDACTED]" in body

    def test_format_counts_per_extension(self, tmp_path):
        _write_txt(tmp_path / "a.txt", "Alpha.")
        _write_txt(tmp_path / "b.md", "Bravo.")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(tmp_path), output_path=str(out), strategy="paragraph", recursive=False)
        assert result.format_counts.get(".txt") == 1
        assert result.format_counts.get(".md") == 1
        assert result.files_processed == 2


class TestIngestDirectory:
    def test_recursive_walks_subdirs(self, tmp_path):
        (tmp_path / "sub").mkdir()
        _write_txt(tmp_path / "top.txt", "Top.")
        _write_txt(tmp_path / "sub" / "deep.txt", "Deep.")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(tmp_path), output_path=str(out), recursive=True, strategy="paragraph")
        assert result.files_processed == 2

    def test_non_recursive_skips_subdirs(self, tmp_path):
        (tmp_path / "sub").mkdir()
        _write_txt(tmp_path / "top.txt", "Top.")
        _write_txt(tmp_path / "sub" / "deep.txt", "Deep.")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(tmp_path), output_path=str(out), recursive=False, strategy="paragraph")
        assert result.files_processed == 1  # only "top.txt"

    def test_skips_unsupported_extensions(self, tmp_path):
        _write_txt(tmp_path / "good.txt", "Good.")
        (tmp_path / "image.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(tmp_path), output_path=str(out), strategy="paragraph")
        # png is not in SUPPORTED_EXTENSIONS so the iterator never sees it
        # — files_skipped counts files we *attempted* and dropped.
        assert result.files_processed == 1


class TestSummary:
    def test_summary_contains_next_step_hint(self, tmp_path):
        src = _write_txt(tmp_path / "doc.txt", "Body.")
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(src), output_path=str(out), strategy="paragraph")
        rendered = summarize_result(result)
        assert "data-audit" in rendered.lower() or "data_audit" in rendered.lower()
        assert "Output JSONL" in rendered


# ---------------------------------------------------------------------------
# Optional-extra extractors — skipped when the dep isn't installed.
# ---------------------------------------------------------------------------


def _has(module_name: str) -> bool:
    try:
        importlib.import_module(module_name)
        return True
    except ImportError:
        return False


@pytest.mark.skipif(not _has("pypdf"), reason="pypdf extra not installed")
class TestPdfExtractor:
    def test_blank_pdf_yields_no_chunks(self, tmp_path):
        # A scanned PDF with no text layer must warn + emit zero chunks
        # rather than crash.
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        pdf_path = tmp_path / "blank.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(pdf_path), output_path=str(out), strategy="paragraph")
        assert result.files_processed == 0  # empty extraction → skipped

    def test_text_bearing_pdf_extracts_chunks(self, tmp_path):
        # Build a real text-bearing PDF by hand — just enough valid PDF to
        # carry a single Tj-encoded string. Avoids pulling in reportlab as a
        # test-only heavy dep just to exercise the happy path.
        pdf_bytes = _hand_built_pdf("Hello world from a real PDF page.")
        pdf_path = tmp_path / "doc.pdf"
        pdf_path.write_bytes(pdf_bytes)

        out = tmp_path / "out.jsonl"
        result = ingest_path(str(pdf_path), output_path=str(out), strategy="paragraph")
        assert result.files_processed == 1
        rows = [json.loads(line) for line in out.read_text(encoding="utf-8").splitlines() if line.strip()]
        assert any("Hello world from a real PDF page" in r["text"] for r in rows)

    def test_encrypted_pdf_raises_clear_error(self, tmp_path):
        # Encrypted PDFs are caught and surfaced; the operator gets an
        # actionable message instead of "extraction failed".
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.encrypt(user_password="secret", owner_password="owner")
        pdf_path = tmp_path / "enc.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # With pii_mask off, the per-file extraction failure surfaces as a
        # warning (not a crash) and the file is skipped.
        out = tmp_path / "out.jsonl"
        result = ingest_path(str(pdf_path), output_path=str(out), strategy="paragraph")
        assert result.files_processed == 0
        assert result.files_skipped >= 1


@pytest.mark.skipif(not _has("docx"), reason="python-docx extra not installed")
class TestDocxExtractor:
    def test_paragraphs_and_tables_round_trip(self, tmp_path):
        # python-docx can author its own input — no third-party fixture
        # needed, no reportlab analogue required. Verifies BOTH the
        # paragraph path AND the row-major table flattening promised by
        # docs/guides/ingestion.md.
        from docx import Document as _Document

        doc = _Document()
        doc.add_paragraph("First paragraph body.")
        doc.add_paragraph("Second paragraph body.")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "left top"
        table.rows[0].cells[1].text = "right top"
        table.rows[1].cells[0].text = "left bot"
        table.rows[1].cells[1].text = "right bot"
        docx_path = tmp_path / "doc.docx"
        doc.save(str(docx_path))

        out = tmp_path / "out.jsonl"
        result = ingest_path(str(docx_path), output_path=str(out), strategy="paragraph")
        assert result.files_processed == 1
        body = out.read_text(encoding="utf-8")
        assert "First paragraph body." in body
        # Row-major flattening: cells joined by " | " on each row
        assert "left top | right top" in body
        assert "left bot | right bot" in body


@pytest.mark.skipif(not _has("ebooklib") or not _has("bs4"), reason="ebooklib/bs4 extras not installed")
class TestEpubExtractor:
    def test_chapter_text_round_trips(self, tmp_path):
        # ebooklib can author its own input. Confirms the option dict
        # we pass to read_epub doesn't break the happy path.
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier("test")
        book.set_title("Test")
        book.set_language("en")
        chapter = epub.EpubHtml(
            title="Chapter",
            file_name="chapter.xhtml",
            content="<html><body><h1>Title</h1><p>Chapter body text here.</p></body></html>",
        )
        book.add_item(chapter)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = [chapter]
        book.toc = (chapter,)
        epub_path = tmp_path / "doc.epub"
        epub.write_epub(str(epub_path), book)

        out = tmp_path / "out.jsonl"
        result = ingest_path(str(epub_path), output_path=str(out), strategy="paragraph")
        assert result.files_processed == 1
        body = out.read_text(encoding="utf-8")
        assert "Chapter body text here." in body


def _hand_built_pdf(text: str) -> bytes:
    """Minimal valid one-page PDF with a single Tj-encoded line of text.

    This is a deliberately tiny, hand-crafted PDF — enough for pypdf's text
    extractor to find a real text layer without dragging in reportlab as a
    test-only dependency. The byte layout follows the PDF 1.4 spec just
    closely enough to load.
    """
    escaped = text.replace("(", "\\(").replace(")", "\\)").replace("\\", "\\\\")
    content_stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    content_obj = (
        b"5 0 obj\n<< /Length "
        + str(len(content_stream)).encode()
        + b" >>\nstream\n"
        + content_stream
        + b"\nendstream\nendobj\n"
    )

    objs = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 5 0 R /Resources << /Font << /F1 4 0 R >> >> >>\nendobj\n",
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
        content_obj,
    ]

    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for obj in objs:
        offsets.append(len(pdf))
        pdf += obj

    xref_offset = len(pdf)
    pdf += b"xref\n0 6\n"
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode()
    pdf += b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
    pdf += str(xref_offset).encode() + b"\n%%EOF"
    return pdf
